import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import date, datetime, timedelta
import psycopg2
import psycopg2.extras
import io
import base64
import json
import hashlib
import re as _re
import sqlite3 as _sqlite3
from pathlib import Path
from PIL import Image as _PILImage

_DB_PATH = Path(__file__).parent / "data" / "cosmetics.db"
USE_LOCAL = _DB_PATH.exists()


class PGConn:
    """psycopg2 래퍼 — sqlite3 conn.execute() 인터페이스 호환"""
    def __init__(self, conn, shared=False):
        self._conn = conn
        self._shared = shared

    def execute(self, sql, params=()):
        sql = sql.replace("?", "%s")
        cur = self._conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(sql, params if params else None)
        return cur

    def commit(self):
        if not self._shared:
            self._conn.commit()

    def close(self):
        if not self._shared:
            self._conn.close()

    def cursor(self):   return self._conn.cursor()


class _PGPool:
    """연결 재사용 풀 — 매 쿼리마다 새 TCP 연결을 맺지 않음"""
    def __init__(self):
        self._url = st.secrets["database"]["url"]
        self._conn = None
        self._connect()

    def _connect(self):
        self._conn = psycopg2.connect(self._url, connect_timeout=10)
        self._conn.autocommit = True

    def conn(self):
        if self._conn is None or self._conn.closed:
            self._connect()
        return self._conn


@st.cache_resource(show_spinner=False)
def _pg_pool():
    return _PGPool()


class SQLiteConn:
    """sqlite3 래퍼 — PGConn과 동일한 인터페이스"""
    def __init__(self, conn):
        self._conn = conn
        self._conn.row_factory = _sqlite3.Row

    def execute(self, sql, params=()):
        sql = _re.sub(r'::\w+', '', sql)
        cur = self._conn.cursor()
        cur.execute(sql, params if params else ())
        return cur

    def commit(self):   self._conn.commit()
    def close(self):    self._conn.close()
    def cursor(self):   return self._conn.cursor()

ADMIN_PW_HASH = hashlib.sha256("0413".encode()).hexdigest()

_FAVICON_PATH = Path(__file__).parent / "data" / "favicon.png"
_SETTINGS_PATH_EARLY = Path(__file__).parent / "data" / "menu_settings.json"

def _load_early_settings():
    if _SETTINGS_PATH_EARLY.exists():
        try:
            d = json.loads(_SETTINGS_PATH_EARLY.read_text(encoding="utf-8"))
            return d.get("site_name", "화장품 유통 관리")
        except Exception:
            pass
    return "화장품 유통 관리"

_page_title = _load_early_settings()
_page_icon  = _PILImage.open(_FAVICON_PATH) if _FAVICON_PATH.exists() else "💄"

st.set_page_config(page_title=_page_title, layout="wide", page_icon=_page_icon,
                   initial_sidebar_state="collapsed")

# ── CSS ────────────────────────────────────────────────────────────────────
def _inject_css(dark: bool = True):
    if dark:
        st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
html,body,[class*="css"]{font-family:'Inter',sans-serif}
.stApp{background:#0b0f1a}
[data-testid="stSidebar"]{background:#0d1117;border-right:1px solid rgba(255,255,255,0.07)}
.kpi-row{display:grid;grid-template-columns:repeat(4,1fr);gap:14px;margin-bottom:20px}
.kpi{background:rgba(255,255,255,0.04);border:1px solid rgba(255,255,255,0.08);
     border-radius:16px;padding:20px 24px;position:relative;overflow:hidden}
.kpi-label{font-size:11px;font-weight:600;color:rgba(255,255,255,0.4);
           letter-spacing:.8px;text-transform:uppercase;margin-bottom:8px}
.kpi-value{font-size:26px;font-weight:700;color:#fff}
.kpi-sub{font-size:12px;margin-top:6px;color:rgba(255,255,255,0.4)}
.card{background:rgba(255,255,255,0.03);border:1px solid rgba(255,255,255,0.07);
      border-radius:16px;padding:20px;margin-bottom:16px}
.card-title{font-size:14px;font-weight:600;color:rgba(255,255,255,0.8);margin-bottom:4px}
.card-sub{font-size:12px;color:rgba(255,255,255,0.35);margin-bottom:14px}
.fifo-box{background:rgba(99,102,241,0.08);border:1px solid rgba(99,102,241,0.3);
          border-radius:10px;padding:12px 16px;margin:10px 0;font-size:13px;color:#a5b4fc}
</style>
""", unsafe_allow_html=True)
    else:
        st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
html,body,[class*="css"]{font-family:'Inter',sans-serif}
.stApp{background:#f5f7fa !important}
[data-testid="stSidebar"]{background:#ffffff !important;border-right:1px solid rgba(0,0,0,0.08) !important}
[data-testid="stAppViewContainer"]{background:#f5f7fa !important}
[data-testid="stHeader"]{background:transparent !important}
.kpi-row{display:grid;grid-template-columns:repeat(4,1fr);gap:14px;margin-bottom:20px}
.kpi{background:#ffffff;border:1px solid rgba(0,0,0,0.08);
     border-radius:16px;padding:20px 24px;position:relative;overflow:hidden;
     box-shadow:0 1px 4px rgba(0,0,0,0.06)}
.kpi-label{font-size:11px;font-weight:600;color:rgba(0,0,0,0.45);
           letter-spacing:.8px;text-transform:uppercase;margin-bottom:8px}
.kpi-value{font-size:26px;font-weight:700;color:#1a1a2e}
.kpi-sub{font-size:12px;margin-top:6px;color:rgba(0,0,0,0.4)}
.card{background:#ffffff;border:1px solid rgba(0,0,0,0.07);
      border-radius:16px;padding:20px;margin-bottom:16px;
      box-shadow:0 1px 4px rgba(0,0,0,0.05)}
.card-title{font-size:14px;font-weight:600;color:rgba(0,0,0,0.8);margin-bottom:4px}
.card-sub{font-size:12px;color:rgba(0,0,0,0.45);margin-bottom:14px}
.fifo-box{background:rgba(99,102,241,0.06);border:1px solid rgba(99,102,241,0.25);
          border-radius:10px;padding:12px 16px;margin:10px 0;font-size:13px;color:#4f46e5}
</style>
""", unsafe_allow_html=True)

# ── DB ─────────────────────────────────────────────────────────────────────
SETTINGS_PATH = Path(__file__).parent / "data" / "menu_settings.json"

DEFAULT_SITE_NAME    = "화장품 유통 관리"
DEFAULT_SITE_CAPTION = "바크로 · 물톡스 통합 관리"
DEFAULT_SITE_ICON    = "💄"

def load_menu_settings(default_menu, default_labels):
    if SETTINGS_PATH.exists():
        try:
            data = json.loads(SETTINGS_PATH.read_text(encoding="utf-8"))
            order   = data.get("order",        default_menu.copy())
            labels  = data.get("labels",       default_labels.copy())
            site    = data.get("site_name",    DEFAULT_SITE_NAME)
            caption = data.get("site_caption", DEFAULT_SITE_CAPTION)
            icon    = data.get("site_icon",    DEFAULT_SITE_ICON)
            theme   = data.get("theme",         "dark")
            for k in default_menu:
                if k not in labels:
                    labels[k] = k
                if k not in order:
                    order.append(k)
            return order, labels, site, caption, icon, theme
        except Exception:
            pass
    return default_menu.copy(), default_labels.copy(), DEFAULT_SITE_NAME, DEFAULT_SITE_CAPTION, DEFAULT_SITE_ICON, "dark"

def save_menu_settings(order, labels, site_name=None, site_caption=None, site_icon=None, theme=None):
    try:
        SETTINGS_PATH.parent.mkdir(exist_ok=True)
        existing = {}
        if SETTINGS_PATH.exists():
            try:
                existing = json.loads(SETTINGS_PATH.read_text(encoding="utf-8"))
            except Exception:
                pass
        existing["order"]        = order
        existing["labels"]       = labels
        existing["site_name"]    = site_name    if site_name    is not None else existing.get("site_name",    DEFAULT_SITE_NAME)
        existing["site_caption"] = site_caption if site_caption is not None else existing.get("site_caption", DEFAULT_SITE_CAPTION)
        existing["site_icon"]    = site_icon    if site_icon    is not None else existing.get("site_icon",    DEFAULT_SITE_ICON)
        existing["theme"]        = theme        if theme        is not None else existing.get("theme",        "dark")
        SETTINGS_PATH.write_text(
            json.dumps(existing, ensure_ascii=False, indent=2),
            encoding="utf-8"
        )
    except Exception:
        pass  # 클라우드 환경에서는 파일 시스템이 읽기 전용일 수 있음

def get_conn():
    if USE_LOCAL:
        return SQLiteConn(_sqlite3.connect(str(_DB_PATH)))
    return PGConn(_pg_pool().conn(), shared=True)

CATEGORY_DISCOUNT = {
    "더데이랩스": 0.20,
    "루트스퀘어": 0.17,
    "비네트워크":  0.10,
}
CATEGORIES = ["", "더데이랩스", "루트스퀘어", "비네트워크"]

def init_db():
    conn = get_conn()
    conn.execute("""CREATE TABLE IF NOT EXISTS products (
        id SERIAL PRIMARY KEY, name TEXT NOT NULL, spec TEXT,
        purchase_price INTEGER DEFAULT 0, sale_price INTEGER DEFAULT 0,
        min_stock INTEGER DEFAULT 10, settlement_price INTEGER DEFAULT 0,
        bundle_with_id INTEGER DEFAULT NULL, product_group TEXT DEFAULT '')""")
    conn.execute("""CREATE TABLE IF NOT EXISTS clients (
        id SERIAL PRIMARY KEY, name TEXT NOT NULL, category TEXT DEFAULT '',
        contact TEXT, phone TEXT, address TEXT,
        credit_limit INTEGER DEFAULT 0, discount_rate REAL DEFAULT 0)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS stock_in (
        id SERIAL PRIMARY KEY, date TEXT NOT NULL, product_id INTEGER NOT NULL,
        quantity INTEGER DEFAULT 0, purchase_price INTEGER DEFAULT 0,
        lot_number TEXT, expiry_date TEXT, remaining_qty INTEGER DEFAULT 0, note TEXT)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS stock_out (
        id SERIAL PRIMARY KEY, date TEXT NOT NULL, client_id INTEGER, product_id INTEGER,
        type TEXT DEFAULT '출고', quantity INTEGER DEFAULT 0, discount_rate REAL DEFAULT 0,
        supply_amount INTEGER DEFAULT 0, vat_amount INTEGER DEFAULT 0,
        total_amount INTEGER DEFAULT 0, fifo_purchase_price INTEGER DEFAULT 0,
        margin_amount INTEGER DEFAULT 0, margin_rate REAL DEFAULT 0, note TEXT)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS payments (
        id SERIAL PRIMARY KEY, date TEXT NOT NULL, client_id INTEGER,
        amount INTEGER DEFAULT 0, linked_out_id INTEGER, note TEXT)""")
    # ── 의료장비 테이블 ──
    conn.execute("""CREATE TABLE IF NOT EXISTS med_products (
        id SERIAL PRIMARY KEY, name TEXT NOT NULL, manufacturer TEXT, spec TEXT)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS med_clients (
        id SERIAL PRIMARY KEY, hospital_name TEXT NOT NULL,
        ceo_name TEXT, contact_person TEXT, phone TEXT, address TEXT,
        credit_limit INTEGER DEFAULT 0)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS med_stock_in (
        id SERIAL PRIMARY KEY, date TEXT NOT NULL, manufacturer TEXT,
        product_name TEXT, price_type TEXT DEFAULT '매입',
        purchase_price INTEGER DEFAULT 0, serial_number TEXT,
        warranty_end TEXT, note TEXT)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS med_stock_out (
        id SERIAL PRIMARY KEY, date TEXT NOT NULL, client_id INTEGER,
        product_id INTEGER, product_name TEXT, manufacturer TEXT,
        serial_number TEXT, is_purchased TEXT DEFAULT '매입',
        payment_method TEXT DEFAULT '현금', sale_price INTEGER DEFAULT 0,
        commission INTEGER DEFAULT 0, supply_amount INTEGER DEFAULT 0,
        vat_amount INTEGER DEFAULT 0, total_amount INTEGER DEFAULT 0,
        drive_file_id TEXT, drive_file_name TEXT, note TEXT)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS med_payments (
        id SERIAL PRIMARY KEY, date TEXT NOT NULL, client_id INTEGER,
        amount INTEGER DEFAULT 0, payment_method TEXT DEFAULT '현금',
        linked_out_id INTEGER, note TEXT)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS med_contracts (
        id SERIAL PRIMARY KEY, client_id INTEGER, linked_out_id INTEGER,
        file_name TEXT, file_type TEXT DEFAULT '납품계약서',
        drive_file_id TEXT, drive_file_url TEXT, uploaded_at TEXT, note TEXT)""")
    conn.execute("""CREATE TABLE IF NOT EXISTS med_settings (
        key TEXT PRIMARY KEY, value TEXT)""")
    conn.commit()
    conn.execute("UPDATE products SET settlement_price=sale_price WHERE settlement_price=0 AND sale_price>0")
    conn.commit()
    conn.close()

def run_sql(query, params=()):
    conn = get_conn()
    try:
        if USE_LOCAL:
            clean = _re.sub(r'::\w+', '', query)
            df = pd.read_sql(clean, conn._conn,
                             params=list(params) if params else None)
        else:
            df = pd.read_sql(query.replace("?", "%s"), conn._conn,
                             params=list(params) if params else None)
    finally:
        conn.close()
    return df

def execute(query, params=()):
    conn = get_conn()
    try:
        conn.execute(query, params)
        conn.commit()
    finally:
        conn.close()
    return None

def fmt(v):    return f"₩{int(v or 0):,}"

# ── FIFO 계산 ──────────────────────────────────────────────────────────────
def get_fifo_info(product_id, qty_needed):
    lots = run_sql("""
        SELECT id, remaining_qty, purchase_price
        FROM stock_in
        WHERE product_id=? AND remaining_qty>0
        ORDER BY date ASC
    """, (product_id,))
    if lots.empty or qty_needed <= 0:
        return 0, [], 0
    available = int(lots["remaining_qty"].sum())
    if qty_needed > available:
        return 0, [], available
    usage, remaining, total_cost = [], qty_needed, 0
    for _, lot in lots.iterrows():
        if remaining <= 0: break
        use = min(remaining, int(lot["remaining_qty"]))
        usage.append({"lot_id": int(lot["id"]), "use_qty": use, "price": int(lot["purchase_price"])})
        total_cost += use * int(lot["purchase_price"])
        remaining -= use
    avg_cost = int(total_cost / qty_needed)
    return avg_cost, usage, available

def apply_fifo(product_id, qty_needed):
    _, usage, _ = get_fifo_info(product_id, qty_needed)
    conn = get_conn()
    for u in usage:
        conn.execute("UPDATE stock_in SET remaining_qty=remaining_qty-? WHERE id=?",
                     (u["use_qty"], u["lot_id"]))
    conn.commit()
    conn.close()

def restore_fifo(out_id):
    conn = get_conn()
    row = conn.execute("SELECT product_id, quantity, type FROM stock_out WHERE id=?", (out_id,)).fetchone()
    if not row: conn.close(); return
    # 단순 복원: 가장 최근 lot에 반환
    conn.execute("""
        UPDATE stock_in SET remaining_qty=remaining_qty+?
        WHERE product_id=? AND id=(SELECT id FROM stock_in WHERE product_id=? ORDER BY date ASC LIMIT 1)
    """, (row["quantity"], row["product_id"], row["product_id"]))
    conn.commit(); conn.close()

# ══════════════════════════════════════════════════════════════════════════
# Supabase Storage 헬퍼
# ══════════════════════════════════════════════════════════════════════════
_STORAGE_BUCKET   = "mso-documents"
_STORAGE_FREE_BYTES = 1 * 1024 * 1024 * 1024  # 1 GB

def _storage_cfg():
    """(base_url, headers) 반환. secrets 미설정 시 (None, None)"""
    try:
        proj_url = st.secrets["supabase"]["url"].rstrip("/")
        key      = st.secrets["supabase"]["service_key"]
        headers  = {"Authorization": f"Bearer {key}", "apikey": key}
        return f"{proj_url}/storage/v1", headers
    except Exception:
        return None, None

def drive_upload(file_bytes, filename, mime_type):
    """파일을 Supabase Storage에 업로드. (path, signed_url) 반환."""
    import requests as _req
    base, hdrs = _storage_cfg()
    if not base:
        st.error("Supabase Storage 설정이 없습니다. secrets.toml을 확인하세요.")
        return None, None
    path = f"{datetime.date.today().isoformat()}_{filename}"
    try:
        r = _req.post(
            f"{base}/object/{_STORAGE_BUCKET}/{path}",
            data=file_bytes,
            headers={**hdrs, "Content-Type": mime_type},
            timeout=60,
        )
        if r.status_code not in (200, 201):
            st.error(f"업로드 오류 ({r.status_code}): {r.text}")
            return None, None
        # 서명된 URL 생성 (1년 유효)
        sign = _req.post(
            f"{base}/object/sign/{_STORAGE_BUCKET}/{path}",
            json={"expiresIn": 31536000},
            headers={**hdrs, "Content-Type": "application/json"},
            timeout=10,
        )
        signed_url = sign.json().get("signedURL", "") if sign.status_code == 200 else ""
        if signed_url and not signed_url.startswith("http"):
            proj_url = st.secrets["supabase"]["url"].rstrip("/")
            signed_url = proj_url + signed_url
        return path, signed_url
    except Exception as e:
        st.error(f"업로드 오류: {e}")
        return None, None

def drive_delete(path):
    """Supabase Storage에서 파일 삭제."""
    import requests as _req
    base, hdrs = _storage_cfg()
    if not base or not path:
        return False
    try:
        r = _req.delete(
            f"{base}/object/{_STORAGE_BUCKET}/{path}",
            headers=hdrs, timeout=10,
        )
        return r.status_code in (200, 204)
    except Exception:
        return False

def drive_download_bytes(path):
    """Supabase Storage에서 파일 바이트 다운로드."""
    import requests as _req
    base, hdrs = _storage_cfg()
    if not base or not path:
        return None
    try:
        r = _req.get(
            f"{base}/object/authenticated/{_STORAGE_BUCKET}/{path}",
            headers=hdrs, timeout=60,
        )
        return r.content if r.status_code == 200 else None
    except Exception:
        return None

def get_storage_usage():
    """사용 중인 바이트 수 반환. 실패 시 -1."""
    import requests as _req
    base, hdrs = _storage_cfg()
    if not base:
        return -1
    try:
        r = _req.post(
            f"{base}/object/list/{_STORAGE_BUCKET}",
            json={"limit": 10000, "offset": 0, "sortBy": {"column": "name", "order": "asc"}},
            headers={**hdrs, "Content-Type": "application/json"},
            timeout=10,
        )
        if r.status_code != 200:
            return -1
        files = r.json()
        return sum(f.get("metadata", {}).get("size", 0) for f in files if isinstance(f.get("metadata"), dict))
    except Exception:
        return -1

def show_storage_warning():
    """용량 경고 배너. 80% 이상 시 표시."""
    used = get_storage_usage()
    if used < 0:
        return
    pct = used / _STORAGE_FREE_BYTES * 100
    used_mb = used / 1024 / 1024
    if pct >= 95:
        st.error(
            f"⛔ 저장 용량 위기! {used_mb:.0f}MB / 1,024MB 사용 중 ({pct:.1f}%)\n\n"
            "**즉시 조치 필요:** Supabase Pro 플랜 업그레이드(월 $25, 100GB) 또는 "
            "Cloudflare R2 무료 플랜(10GB) 전환을 권장합니다."
        )
    elif pct >= 80:
        st.warning(
            f"⚠️ 저장 용량 {pct:.1f}% 사용 중 ({used_mb:.0f}MB / 1,024MB)\n\n"
            "**용량 확장 옵션:** ① Supabase Pro (월 $25, 100GB) "
            "② Cloudflare R2 무료 10GB 전환"
        )

# ══════════════════════════════════════════════════════════════════════════
# 대시보드
# ══════════════════════════════════════════════════════════════════════════
def page_dashboard():

    today        = date.today()
    today_str    = today.strftime("%Y-%m-%d")
    # 캘린더 기준: 이번 주 월요일
    week_start   = (today - timedelta(days=today.weekday())).strftime("%Y-%m-%d")
    # 이번 달 1일
    month_start  = today.replace(day=1).strftime("%Y-%m-%d")
    # 올해 1월 1일
    year_start   = today.replace(month=1, day=1).strftime("%Y-%m-%d")
    this_month   = today.strftime("%Y-%m")
    this_year    = today.strftime("%Y")

    _s = run_sql("""
        SELECT
            COALESCE(SUM(CASE WHEN date=? THEN total_amount END),0)              as today,
            COALESCE(SUM(CASE WHEN date BETWEEN ? AND ? THEN total_amount END),0) as week,
            COALESCE(SUM(CASE WHEN date BETWEEN ? AND ? THEN total_amount END),0) as month,
            COALESCE(SUM(CASE WHEN date BETWEEN ? AND ? THEN total_amount END),0) as year
        FROM stock_out WHERE type='출고'
    """, (today_str, week_start, today_str, month_start, today_str, year_start, today_str)).iloc[0]
    today_sales  = _s["today"]
    week_sales   = _s["week"]
    month_sales  = _s["month"]
    year_sales   = _s["year"]

    total_recv   = run_sql("""
        SELECT
            (SELECT COALESCE(SUM(total_amount),0) FROM stock_out WHERE type='출고')
            - (SELECT COALESCE(SUM(amount),0) FROM payments) as v
    """).iloc[0]["v"]
    month_margin = run_sql("SELECT COALESCE(AVG(margin_rate),0) as v FROM stock_out WHERE date LIKE ? AND type='출고'", (f"{this_month}%",)).iloc[0]["v"]
    low_stock_items = run_sql("""
        SELECT p.name as 제품명, COALESCE(s.rq,0) as 현재고, p.min_stock as 최소재고
        FROM products p
        LEFT JOIN (SELECT product_id, SUM(remaining_qty) as rq FROM stock_in GROUP BY product_id) s ON p.id=s.product_id
        WHERE COALESCE(s.rq,0) < p.min_stock
        ORDER BY COALESCE(s.rq,0) ASC
    """)
    low_stock_cnt = len(low_stock_items)

    # 매출 리포트
    st.markdown(f"""
    <div class="kpi-row">
      <div class="kpi"><div class="kpi-label">오늘 매출</div><div class="kpi-value">{fmt(today_sales)}</div><div class="kpi-sub">{today_str}</div></div>
      <div class="kpi"><div class="kpi-label">주간 매출</div><div class="kpi-value">{fmt(week_sales)}</div><div class="kpi-sub">{week_start} ~ {today_str}</div></div>
      <div class="kpi"><div class="kpi-label">월간 매출</div><div class="kpi-value">{fmt(month_sales)}</div><div class="kpi-sub">{month_start} ~ {today_str}</div></div>
      <div class="kpi"><div class="kpi-label">연간 매출</div><div class="kpi-value">{fmt(year_sales)}</div><div class="kpi-sub">{year_start} ~ {today_str}</div></div>
    </div>
    """, unsafe_allow_html=True)

    # 주요 현황
    st.markdown(f"""
    <div class="kpi-row">
      <div class="kpi"><div class="kpi-label">미수금 합계</div><div class="kpi-value">{fmt(total_recv)}</div><div class="kpi-sub">전체 누적</div></div>
      <div class="kpi"><div class="kpi-label">이번달 평균 마진율</div><div class="kpi-value">{month_margin:.1f}%</div><div class="kpi-sub">{this_month}</div></div>
      <div class="kpi"><div class="kpi-label">재고 부족 품목</div><div class="kpi-value">{low_stock_cnt}개</div><div class="kpi-sub">최소 재고 미달 ↓ 클릭</div></div>
    </div>
    """, unsafe_allow_html=True)

    if low_stock_cnt > 0:
        with st.expander(f"⚠️ 재고 부족 품목 {low_stock_cnt}개 — 상세 보기", expanded=False):
            st.dataframe(low_stock_items, use_container_width=True, hide_index=True)

    # 월별 / 연도별 그래프 나란히
    gcol1, gcol2 = st.columns(2, gap="medium")

    with gcol1:
        st.markdown('<div class="card"><div class="card-title">월별 매출 추이</div><div class="card-sub">출고 기준</div>', unsafe_allow_html=True)
        monthly = run_sql("""
            SELECT substr(date,1,7) as 월, SUM(total_amount) as 매출, SUM(margin_amount) as 마진
            FROM stock_out WHERE type='출고' GROUP BY 월 ORDER BY 월
        """)
        if not monthly.empty:
            fig = go.Figure()
            fig.add_trace(go.Bar(x=monthly["월"], y=monthly["매출"], name="매출", marker_color="#6366f1"))
            fig.add_trace(go.Scatter(x=monthly["월"], y=monthly["마진"], name="마진", line=dict(color="#34d399", width=2), yaxis="y2"))
            fig.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                              yaxis2=dict(overlaying="y", side="right", gridcolor="rgba(0,0,0,0)"),
                              xaxis=dict(gridcolor="rgba(255,255,255,0.05)", tickfont=dict(color="rgba(255,255,255,0.4)")),
                              yaxis=dict(gridcolor="rgba(255,255,255,0.05)", tickfont=dict(color="rgba(255,255,255,0.4)")),
                              legend=dict(font=dict(color="rgba(255,255,255,0.6)")),
                              margin=dict(l=0,r=0,t=10,b=0))
            st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
        else:
            st.info("출고 데이터가 없습니다.")
        st.markdown("</div>", unsafe_allow_html=True)

    with gcol2:
        st.markdown('<div class="card"><div class="card-title">연도별 매출 추이</div><div class="card-sub">출고 기준</div>', unsafe_allow_html=True)
        yearly = run_sql("""
            SELECT substr(date,1,4) as 연도, SUM(total_amount) as 매출, SUM(margin_amount) as 마진
            FROM stock_out WHERE type='출고' GROUP BY 연도 ORDER BY 연도
        """)
        if not yearly.empty:
            fig2 = go.Figure()
            fig2.add_trace(go.Bar(x=yearly["연도"], y=yearly["매출"], name="매출", marker_color="#818cf8"))
            fig2.add_trace(go.Scatter(x=yearly["연도"], y=yearly["마진"], name="마진", line=dict(color="#34d399", width=2), yaxis="y2"))
            fig2.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                               yaxis2=dict(overlaying="y", side="right", gridcolor="rgba(0,0,0,0)"),
                               xaxis=dict(gridcolor="rgba(255,255,255,0.05)", tickfont=dict(color="rgba(255,255,255,0.4)")),
                               yaxis=dict(gridcolor="rgba(255,255,255,0.05)", tickfont=dict(color="rgba(255,255,255,0.4)")),
                               legend=dict(font=dict(color="rgba(255,255,255,0.6)")),
                               margin=dict(l=0,r=0,t=10,b=0))
            st.plotly_chart(fig2, use_container_width=True, config={"displayModeBar": False})
        else:
            st.info("출고 데이터가 없습니다.")
        st.markdown("</div>", unsafe_allow_html=True)

    # 미수금 TOP 5
    st.markdown('<div class="card"><div class="card-title">미수금 TOP 5 거래처</div><div class="card-sub">높은 순</div>', unsafe_allow_html=True)
    top_recv = run_sql("""
        SELECT name, 미수금 FROM (
            SELECT c.name,
                COALESCE(s.total_sales,0) - COALESCE(p.total_paid,0) as 미수금
            FROM clients c
            LEFT JOIN (
                SELECT client_id, SUM(total_amount) as total_sales
                FROM stock_out WHERE type='출고' GROUP BY client_id
            ) s ON c.id=s.client_id
            LEFT JOIN (
                SELECT client_id, SUM(amount) as total_paid
                FROM payments GROUP BY client_id
            ) p ON c.id=p.client_id
        ) WHERE 미수금>0 ORDER BY 미수금 DESC LIMIT 5
    """)
    if not top_recv.empty:
        top_recv["미수금"] = top_recv["미수금"].apply(fmt)
        st.dataframe(top_recv, use_container_width=True, hide_index=True)
    else:
        st.success("미수금 없음")
    st.markdown("</div>", unsafe_allow_html=True)

    # ── 엑셀 다운로드 ──────────────────────────────────────────────
    st.markdown('<div class="card"><div class="card-title">📥 리포트 엑셀 다운로드</div>', unsafe_allow_html=True)

    col_period1, col_period2, col_dl = st.columns([2, 2, 2])
    xl_from = col_period1.date_input("시작일", value=date(today.year, 1, 1), key="xl_from")
    xl_to   = col_period2.date_input("종료일", value=today, key="xl_to")

    if col_dl.button("📊 엑셀 생성", type="primary", use_container_width=True, key="gen_excel"):
        import io
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        frm, to_ = str(xl_from), str(xl_to)

        # ── 데이터 조회 ──
        df_out = run_sql("""
            SELECT o.date as 날짜, c.name as 거래처, p.name as 품목,
                   o.quantity as 수량,
                   ROUND((o.discount_rate*100)::NUMERIC,1) as 할인율PCT,
                   COALESCE(o.fifo_purchase_price,0)*o.quantity as 공급가액,
                   o.total_amount as 납품금액,
                   o.margin_amount as 마진액,
                   ROUND(o.margin_rate::NUMERIC,2) as 마진율PCT,
                   COALESCE(o.note,'') as 메모
            FROM stock_out o
            JOIN clients c ON o.client_id=c.id
            JOIN products p ON o.product_id=p.id
            WHERE o.type='출고' AND o.date BETWEEN ? AND ?
            ORDER BY o.date, c.name
        """, (frm, to_))

        df_client = run_sql("""
            SELECT c.name as 거래처,
                   COALESCE(s.ts,0) as 납품금액합계,
                   COALESCE(p.tp,0) as 수금액합계,
                   COALESCE(s.ts,0)-COALESCE(p.tp,0) as 미수금,
                   COALESCE(s.mg,0) as 마진액합계,
                   CASE WHEN COALESCE(s.ts,0)>0
                        THEN ROUND((COALESCE(s.mg,0)*100.0/COALESCE(s.ts,0))::NUMERIC,2)
                        ELSE 0 END as 평균마진율PCT
            FROM clients c
            LEFT JOIN (
                SELECT client_id,
                       SUM(total_amount) as ts,
                       SUM(margin_amount) as mg
                FROM stock_out
                WHERE type='출고' AND date BETWEEN ? AND ?
                GROUP BY client_id
            ) s ON c.id=s.client_id
            LEFT JOIN (
                SELECT client_id, SUM(amount) as tp
                FROM payments WHERE date BETWEEN ? AND ?
                GROUP BY client_id
            ) p ON c.id=p.client_id
            WHERE COALESCE(s.ts,0)>0
            ORDER BY 납품금액합계 DESC
        """, (frm, to_, frm, to_))

        df_monthly = run_sql("""
            SELECT substr(date,1,7) as 월,
                   SUM(total_amount) as 납품금액,
                   SUM(margin_amount) as 마진액,
                   ROUND(AVG(margin_rate)::NUMERIC,2) as 평균마진율PCT
            FROM stock_out WHERE type='출고' AND date BETWEEN ? AND ?
            GROUP BY 월 ORDER BY 월
        """, (frm, to_))

        df_stock = run_sql("""
            SELECT p.name as 제품명, p.spec as 규격,
                   p.purchase_price as 공급가,
                   p.settlement_price as 납품가,
                   COALESCE(SUM(s.remaining_qty),0) as 현재고,
                   p.min_stock as 최소재고
            FROM products p
            LEFT JOIN stock_in s ON p.id=s.product_id
            GROUP BY p.id ORDER BY p.name
        """)

        df_in = run_sql("""
            SELECT s.date as 날짜, p.name as 품목,
                   s.quantity as 수량,
                   s.purchase_price as 단가,
                   s.quantity * s.purchase_price as 합계VAT
            FROM stock_in s JOIN products p ON s.product_id=p.id
            WHERE s.date BETWEEN ? AND ?
            ORDER BY s.date, p.name
        """, (frm, to_))

        # ── 워크북 생성 ──
        wb = Workbook()

        # 스타일 정의
        hdr_font  = Font(bold=True, color="FFFFFF", size=10)
        hdr_fill  = PatternFill("solid", fgColor="4F46E5")
        sub_fill  = PatternFill("solid", fgColor="E0E7FF")
        total_fill= PatternFill("solid", fgColor="C7D2FE")
        center    = Alignment(horizontal="center", vertical="center")
        right_al  = Alignment(horizontal="right")
        thin      = Side(style="thin", color="CCCCCC")
        border    = Border(left=thin, right=thin, top=thin, bottom=thin)

        def style_header(ws, row, cols):
            for c in range(1, cols+1):
                cell = ws.cell(row=row, column=c)
                cell.font = hdr_font; cell.fill = hdr_fill
                cell.alignment = center; cell.border = border

        def style_row(ws, row, cols, fill=None):
            for c in range(1, cols+1):
                cell = ws.cell(row=row, column=c)
                if fill: cell.fill = fill
                cell.border = border

        def auto_width(ws):
            for col in ws.columns:
                max_len = max((len(str(cell.value or "")) for cell in col), default=0)
                ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 4, 30)

        # ── 시트1: 입고 상세 ──
        ws_in = wb.active; ws_in.title = "입고 상세"
        ws_in.append(["날짜","품목","수량","단가","입고가액(공급가액)","부가세","합계"])
        style_header(ws_in, 1, 7)
        tot_in_공급 = tot_in_vat = tot_in_합계 = 0
        for _, r in df_in.iterrows():
            합계  = int(r["합계VAT"])
            공급가액 = round(합계 / 1.1)
            부가세   = 합계 - 공급가액
            ws_in.append([r["날짜"], r["품목"], int(r["수량"]), int(r["단가"]),
                          공급가액, 부가세, 합계])
            style_row(ws_in, ws_in.max_row, 7)
            tot_in_공급 += 공급가액; tot_in_vat += 부가세; tot_in_합계 += 합계
        if not df_in.empty:
            tr = ws_in.max_row + 1
            ws_in.cell(tr,1,"합계"); ws_in.cell(tr,3,int(df_in["수량"].sum()))
            ws_in.cell(tr,5,tot_in_공급); ws_in.cell(tr,6,tot_in_vat); ws_in.cell(tr,7,tot_in_합계)
            style_row(ws_in, tr, 7, total_fill)
            ws_in.cell(tr,1).font = Font(bold=True)
        auto_width(ws_in)

        # ── 시트2: 출고 상세 ──
        ws1 = wb.create_sheet("출고 상세")
        ws1.append(["날짜","거래처","품목","수량","할인율(%)",
                    "출고가액(공급가액)","부가세","합계(납품금액)",
                    "매입원가합계","마진액","마진율(%)","메모"])
        style_header(ws1, 1, 12)
        tot_공급 = tot_vat = tot_합계 = tot_매입 = tot_마진 = 0
        for _, r in df_out.iterrows():
            합계    = int(r["납품금액"])
            출고가액 = round(합계 / 1.1)
            부가세   = 합계 - 출고가액
            매입원가 = int(r["공급가액"])
            마진액   = int(r["마진액"])
            마진율   = float(r["마진율PCT"])
            ws1.append([r["날짜"], r["거래처"], r["품목"], int(r["수량"]),
                        float(r["할인율PCT"]),
                        출고가액, 부가세, 합계,
                        매입원가, 마진액, 마진율, r["메모"]])
            style_row(ws1, ws1.max_row, 12)
            tot_공급 += 출고가액; tot_vat += 부가세; tot_합계 += 합계
            tot_매입 += 매입원가; tot_마진 += 마진액
        if not df_out.empty:
            tr = ws1.max_row + 1
            ws1.cell(tr,1,"합계"); ws1.cell(tr,4,int(df_out["수량"].sum()))
            ws1.cell(tr,6,tot_공급); ws1.cell(tr,7,tot_vat); ws1.cell(tr,8,tot_합계)
            ws1.cell(tr,9,tot_매입); ws1.cell(tr,10,tot_마진)
            ws1.cell(tr,11, round(tot_마진/tot_합계*100, 2) if tot_합계 else 0)
            style_row(ws1, tr, 12, total_fill)
            ws1.cell(tr,1).font = Font(bold=True)
        auto_width(ws1)

        # ── 시트3: 거래처별 요약 ──
        ws2 = wb.create_sheet("거래처별 요약")
        ws2.append(["거래처","납품금액합계","수금액합계","미수금","마진액합계","평균마진율(%)"])
        style_header(ws2, 1, 6)
        for _, r in df_client.iterrows():
            ws2.append([r["거래처"], int(r["납품금액합계"]), int(r["수금액합계"]),
                        int(r["미수금"]), int(r["마진액합계"]), float(r["평균마진율PCT"])])
            style_row(ws2, ws2.max_row, 6)
        if not df_client.empty:
            tr = ws2.max_row + 1
            ws2.cell(tr,1,"합계")
            for ci, col in enumerate(["납품금액합계","수금액합계","미수금","마진액합계"], start=2):
                ws2.cell(tr, ci, int(df_client[col].sum()))
            style_row(ws2, tr, 6, total_fill)
            ws2.cell(tr,1).font = Font(bold=True)
        auto_width(ws2)

        # ── 시트4: 월별 매출 ──
        ws3 = wb.create_sheet("월별 매출")
        ws3.append(["월","납품금액","마진액","평균마진율(%)"])
        style_header(ws3, 1, 4)
        for _, r in df_monthly.iterrows():
            ws3.append([r["월"], int(r["납품금액"]), int(r["마진액"]), float(r["평균마진율PCT"])])
            style_row(ws3, ws3.max_row, 4)
        if not df_monthly.empty:
            tr = ws3.max_row + 1
            ws3.cell(tr,1,"합계")
            ws3.cell(tr,2,int(df_monthly["납품금액"].sum()))
            ws3.cell(tr,3,int(df_monthly["마진액"].sum()))
            style_row(ws3, tr, 4, total_fill)
            ws3.cell(tr,1).font = Font(bold=True)
        auto_width(ws3)

        # ── 시트5: 재고 현황 ──
        ws4 = wb.create_sheet("재고 현황")
        ws4.append(["제품명","규격","공급가","납품가","현재고","최소재고","재고상태"])
        style_header(ws4, 1, 7)
        for _, r in df_stock.iterrows():
            status = "⚠ 부족" if int(r["현재고"]) < int(r["최소재고"]) else "정상"
            ws4.append([r["제품명"], r["규격"] or "",
                        int(r["공급가"]), int(r["납품가"]),
                        int(r["현재고"]), int(r["최소재고"]), status])
            fill = PatternFill("solid", fgColor="FEE2E2") if status == "⚠ 부족" else None
            style_row(ws4, ws4.max_row, 7, fill)
        auto_width(ws4)

        # ── 저장 ──
        buf = io.BytesIO()
        wb.save(buf); buf.seek(0)
        fname = f"SC-MSO_리포트_{frm}_{to_}.xlsx"
        st.download_button("⬇ 엑셀 다운로드", data=buf,
                           file_name=fname,
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           use_container_width=True)

    st.markdown("</div>", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════
# 매입/매출 입력
# ══════════════════════════════════════════════════════════════════════════
def page_stock_entry():
    tab1, tab2, tab3 = st.tabs(["📥 입고 등록", "📤 출고 등록", "✏️ 출고 수정"])

    products = run_sql("SELECT id, name, purchase_price, sale_price, settlement_price FROM products")
    clients  = run_sql("SELECT id, name, category, discount_rate FROM clients")

    # ── 입고 ──
    with tab1:
        if "in_cart" not in st.session_state:
            st.session_state.in_cart = []
        st.markdown('<div class="card">', unsafe_allow_html=True)
        if products.empty:
            st.warning("먼저 제품을 등록해주세요.")
        else:
            in_date = st.date_input("입고일", value=date.today(), key="in_date")
            st.divider()
            st.markdown("**품목 추가**")
            c1, c2 = st.columns([3, 1])
            prod_sel = c1.selectbox("품목 선택", products["name"].tolist(), key="in_prod_sel")
            prod_row = products[products["name"] == prod_sel].iloc[0]
            ia, ib = st.columns(2)
            qty_add   = ia.number_input("수량", min_value=1, value=1, key="in_qty_add")
            price_add = ib.number_input(
                "입고가", min_value=0, step=1000,
                value=int(prod_row["purchase_price"]),
                key=f"in_price_{prod_sel}"
            )
            if st.button("➕ 목록에 추가", key="in_add_btn", use_container_width=True):
                st.session_state.in_cart.append({
                    "product_id": int(prod_row["id"]), "품목": prod_sel,
                    "수량": qty_add, "입고가": price_add,
                    "소계": price_add * qty_add,
                }); st.rerun()
            if st.session_state.in_cart:
                st.divider()
                st.markdown("**입고 목록**")
                disp_in = pd.DataFrame(st.session_state.in_cart)[["품목", "수량", "입고가", "소계"]].copy()
                disp_in.index = range(1, len(disp_in) + 1)
                disp_in["입고가"] = disp_in["입고가"].apply(fmt)
                disp_in["소계"]   = disp_in["소계"].apply(fmt)
                st.dataframe(disp_in, use_container_width=True)
                dc1, dc2 = st.columns([1, 3])
                del_i = dc1.number_input("삭제할 행 번호", min_value=1,
                                          max_value=len(st.session_state.in_cart),
                                          value=1, step=1, key="in_del_idx")
                if dc2.button("선택 행 삭제", key="in_del_btn"):
                    st.session_state.in_cart.pop(del_i - 1); st.rerun()
                st.metric("합계 매입금액", fmt(sum(i["소계"] for i in st.session_state.in_cart)))
                in_note = st.text_input("공통 메모", key="in_note_final")
                sb1, sb2 = st.columns(2)
                if sb1.button("✅ 입고 등록 완료", type="primary", use_container_width=True):
                    cnt = len(st.session_state.in_cart)
                    conn = get_conn()
                    for item in st.session_state.in_cart:
                        conn.execute("""INSERT INTO stock_in
                            (date,product_id,quantity,purchase_price,remaining_qty,note)
                            VALUES (?,?,?,?,?,?)""",
                            (str(in_date), item["product_id"], item["수량"],
                             item["입고가"], item["수량"], in_note))
                    conn.commit(); conn.close()
                    st.session_state.in_cart = []
                    st.success(f"입고 등록 완료 — {cnt}개 품목"); st.rerun()
                if sb2.button("🗑 전체 초기화", use_container_width=True, key="in_clear"):
                    st.session_state.in_cart = []; st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
        recent_in = run_sql("""SELECT s.id, s.date as 날짜, p.name as 품목,
                                      s.quantity as 수량, s.purchase_price as 입고가,
                                      (s.quantity * s.purchase_price) as 입고합계,
                                      s.remaining_qty as 잔여수량
                               FROM stock_in s JOIN products p ON s.product_id=p.id
                               ORDER BY s.date DESC""")
        if not recent_in.empty:
            st.subheader("전체 입고 내역")
            warn_in = recent_in[recent_in["잔여수량"] < recent_in["수량"]]
            for _, wr in warn_in.iterrows():
                used = int(wr["수량"]) - int(wr["잔여수량"])
                st.caption(f"⚠️ {wr['날짜']} | {wr['품목']}: {used}개 이미 출고 사용됨 — 삭제 주의")
            display_in = recent_in.copy()
            display_in.insert(0, "삭제", False)
            edited_in = st.data_editor(
                display_in,
                use_container_width=True, hide_index=True, num_rows="fixed",
                disabled=["날짜", "품목", "수량", "입고가", "입고합계", "잔여수량"],
                column_config={
                    "id":       None,
                    "삭제":     st.column_config.CheckboxColumn("삭제", width="small"),
                    "입고가":   st.column_config.NumberColumn("입고가",   format="₩%d"),
                    "입고합계": st.column_config.NumberColumn("입고합계", format="₩%d"),
                },
                key="editor_in_hist",
            )
            if st.button("🗑 삭제 적용", key="apply_del_in", type="primary"):
                to_delete = edited_in[edited_in["삭제"] == True]
                if not to_delete.empty:
                    conn = get_conn()
                    for did in to_delete["id"].astype(int):
                        conn.execute("DELETE FROM stock_in WHERE id=?", (did,))
                    conn.commit(); conn.close()
                    st.success(f"{len(to_delete)}건 삭제 완료")
                    st.rerun()
                else:
                    st.info("삭제할 항목의 '삭제' 체크박스를 선택해주세요.")

    # ── 출고 ──
    with tab2:
        if "out_cart" not in st.session_state:
            st.session_state.out_cart = []
        if "out_last_client" not in st.session_state:
            st.session_state.out_last_client = None
        st.markdown('<div class="card">', unsafe_allow_html=True)
        if clients.empty or products.empty:
            st.warning("거래처와 제품을 먼저 등록해주세요.")
        else:
            oc1, oc2 = st.columns(2)
            out_type = oc1.radio("유형", ["출고", "반품"], horizontal=True)
            out_date = oc2.date_input("날짜", value=date.today(), key="out_date")
            client_sel = oc1.selectbox("거래처", clients["name"].tolist(), key="out_client")
            client_row = clients[clients["name"] == client_sel].iloc[0]
            discount_rate = float(client_row["discount_rate"])
            cat_label = f"[{client_row['category']}] " if client_row.get("category") else ""
            oc2.info(f"{cat_label}할인율: **{discount_rate*100:.1f}%**")
            if st.session_state.out_last_client != client_sel:
                st.session_state.out_cart = []
                st.session_state.out_last_client = client_sel
            st.divider()
            st.markdown("**품목 추가**")
            prod_sel2 = st.selectbox("품목 선택", products["name"].tolist(), key="out_prod_sel")
            prod_row2 = products[products["name"] == prod_sel2].iloc[0]
            base_price2 = int(prod_row2["settlement_price"] or prod_row2["sale_price"])
            avail_db = int(run_sql(
                "SELECT COALESCE(SUM(remaining_qty),0) as v FROM stock_in WHERE product_id=?",
                (int(prod_row2["id"]),)).iloc[0]["v"])
            carted_qty = sum(i["수량"] for i in st.session_state.out_cart
                             if i["product_id"] == int(prod_row2["id"]))
            remaining_avail = max(0, avail_db - carted_qty)
            pc1, pc2 = st.columns([3, 1])
            qty_add2 = pc1.number_input("수량", min_value=1,
                                         max_value=max(1, remaining_avail),
                                         value=1, key="out_qty_add")
            pc2.metric("가용 재고", f"{remaining_avail}개")
            if remaining_avail <= 0:
                st.error("⚠️ 재고가 없습니다.")
            else:
                if carted_qty > 0:
                    st.caption(f"이미 {carted_qty}개 목록 추가됨 — FIFO 단가는 등록 시 재계산")
                avg_cost2, _, _ = get_fifo_info(int(prod_row2["id"]), qty_add2)
                # 납품가(VAT포함) × (1-할인율) = 고객 납품단가(VAT포함)
                disc_unit_incl = int(base_price2 * (1 - discount_rate))
                납품금액2     = disc_unit_incl * qty_add2          # 납품금액(VAT포함, 판매가 합계)
                공급가액2     = avg_cost2 * qty_add2               # 공급가액(VAT포함, 매입원가 합계)
                margin_total2 = 납품금액2 - 공급가액2
                margin_rate2  = (margin_total2 / 납품금액2 * 100) if 납품금액2 > 0 else 0

                # 세트 구성품 확인
                bundle_row2 = run_sql(
                    "SELECT id, name FROM products WHERE id=(SELECT bundle_with_id FROM products WHERE id=?)",
                    (int(prod_row2["id"]),)
                )
                has_bundle2 = not bundle_row2.empty
                bundle_avail2 = 0
                if has_bundle2:
                    bundle_avail2 = int(run_sql(
                        "SELECT COALESCE(SUM(remaining_qty),0) as v FROM stock_in WHERE product_id=?",
                        (int(bundle_row2.iloc[0]["id"]),)).iloc[0]["v"])

                box_text = (f'납품단가(VAT포함): <b>{fmt(disc_unit_incl)}</b> | '
                            f'공급가액(VAT포함): <b>{fmt(avg_cost2)}</b> | '
                            f'예상 마진율: <b>{margin_rate2:.1f}%</b>')
                if has_bundle2:
                    box_text += f' | 🔗 세트: <b>{bundle_row2.iloc[0]["name"]}</b> 자동 추가'
                st.markdown(f'<div class="fifo-box">{box_text}</div>', unsafe_allow_html=True)

                if has_bundle2 and bundle_avail2 < qty_add2:
                    st.warning(f"⚠️ 세트 구성품 '{bundle_row2.iloc[0]['name']}' 재고 부족 (가용: {bundle_avail2}개)")
                if margin_rate2 < 0:
                    st.error(f"🚨 마진율 음수 ({margin_rate2:.1f}%) — 확인 필요!")
                elif margin_rate2 < 10:
                    st.warning(f"⚠️ 마진율 낮음 ({margin_rate2:.1f}%)")
                if st.button("➕ 목록에 추가", key="out_add_btn", use_container_width=True):
                    st.session_state.out_cart.append({
                        "product_id": int(prod_row2["id"]), "품목": prod_sel2,
                        "수량": qty_add2, "FIFO매입가": avg_cost2,
                        "납품단가": disc_unit_incl,
                        "공급가액": 공급가액2,   # 매입원가 합계 (표시용)
                        "납품금액": 납품금액2,   # 판매가 합계 (DB supply_amount 저장)
                        "VAT": 0, "합계": 납품금액2,
                        "마진액": margin_total2, "마진율": margin_rate2,
                        "discount_rate": discount_rate,
                    })
                    # 세트 구성품 자동 추가
                    if has_bundle2:
                        b_id   = int(bundle_row2.iloc[0]["id"])
                        b_name = bundle_row2.iloc[0]["name"]
                        b_cost, _, _ = get_fifo_info(b_id, qty_add2)
                        st.session_state.out_cart.append({
                            "product_id": b_id,
                            "품목": f"{b_name} (세트)",
                            "수량": qty_add2, "FIFO매입가": b_cost,
                            "납품단가": 0,
                            "공급가액": b_cost * qty_add2,
                            "납품금액": 0, "VAT": 0, "합계": 0,
                            "마진액": -(b_cost * qty_add2),
                            "마진율": 0.0,
                            "discount_rate": 0.0,
                        })
                    st.rerun()
            if st.session_state.out_cart:
                st.divider()
                st.markdown("**출고 목록**")
                disp_out = pd.DataFrame(st.session_state.out_cart)[
                    ["품목", "수량", "납품단가", "공급가액", "마진율"]].copy()
                disp_out.index = range(1, len(disp_out) + 1)
                disp_out["납품단가"] = disp_out["납품단가"].apply(fmt)
                disp_out["공급가액"] = disp_out["공급가액"].apply(fmt)
                disp_out["마진율"]   = disp_out["마진율"].apply(lambda x: f"{x:.1f}%")
                st.dataframe(disp_out, use_container_width=True)
                tot_공급가액 = sum(i["공급가액"] for i in st.session_state.out_cart)
                tot_납품금액 = sum(i["납품금액"] for i in st.session_state.out_cart)
                tot_margin   = sum(i["마진액"]   for i in st.session_state.out_cart)
                avg_mr = (tot_margin / tot_납품금액 * 100) if tot_납품금액 > 0 else 0
                sm1, sm2, sm3 = st.columns(3)
                sm1.metric("공급가액(VAT포함)", fmt(tot_공급가액))
                sm2.metric("납품금액(VAT포함)", fmt(tot_납품금액))
                sm3.metric("평균 마진율", f"{avg_mr:.1f}%")
                dc3, dc4 = st.columns([1, 3])
                del_j = dc3.number_input("삭제할 행 번호", min_value=1,
                                          max_value=len(st.session_state.out_cart),
                                          value=1, step=1, key="out_del_idx")
                if dc4.button("선택 행 삭제", key="out_del_btn"):
                    st.session_state.out_cart.pop(del_j - 1); st.rerun()
                out_note2 = st.text_input("메모", key="out_note_final")
                sb3, sb4 = st.columns(2)
                if sb3.button("✅ 출고 등록 완료", type="primary", use_container_width=True):
                    cnt2 = len(st.session_state.out_cart)
                    tot2 = sum(i["합계"] for i in st.session_state.out_cart)
                    conn = get_conn()
                    for item in st.session_state.out_cart:
                        conn.execute("""INSERT INTO stock_out
                            (date,client_id,product_id,type,quantity,discount_rate,
                             supply_amount,vat_amount,total_amount,fifo_purchase_price,
                             margin_amount,margin_rate,note)
                            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                            (str(out_date), int(client_row["id"]), item["product_id"], out_type,
                             item["수량"], item["discount_rate"], item["납품금액"],
                             item["VAT"], item["합계"], item["FIFO매입가"],
                             item["마진액"], item["마진율"], out_note2))
                    conn.commit(); conn.close()
                    if out_type == "출고":
                        for item in st.session_state.out_cart:
                            apply_fifo(item["product_id"], item["수량"])
                    st.session_state.out_cart = []
                    st.success(f"출고 등록 완료 — {cnt2}개 품목 ({fmt(tot2)})"); st.rerun()
                if sb4.button("🗑 전체 초기화", use_container_width=True, key="out_clear"):
                    st.session_state.out_cart = []; st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)
        recent_out = run_sql("""SELECT o.id, o.date as 날짜, c.name as 거래처, p.name as 품목,
                                       o.type as 유형, o.quantity as 수량,
                                       o.total_amount as 납품금액, o.margin_rate as 마진율
                                FROM stock_out o
                                JOIN clients c ON o.client_id=c.id
                                JOIN products p ON o.product_id=p.id
                                ORDER BY o.date DESC""")
        if not recent_out.empty:
            st.subheader("전체 출고 내역")
            st.caption("삭제할 행의 '삭제' 체크박스 선택 → '삭제 적용' 클릭")
            display_out = recent_out.copy()
            display_out.insert(0, "삭제", False)
            edited_out = st.data_editor(
                display_out,
                use_container_width=True, hide_index=True, num_rows="fixed",
                disabled=["날짜", "거래처", "품목", "유형", "수량", "납품금액", "마진율"],
                column_config={
                    "id":       None,
                    "삭제":     st.column_config.CheckboxColumn("삭제", width="small"),
                    "납품금액": st.column_config.NumberColumn("납품금액", format="₩%d"),
                    "마진율":   st.column_config.NumberColumn("마진율",   format="%.1f%%"),
                },
                key="editor_out_hist",
            )
            if st.button("🗑 삭제 적용", key="apply_del_out", type="primary"):
                to_delete = edited_out[edited_out["삭제"] == True]
                if not to_delete.empty:
                    del_ids = to_delete["id"].astype(int).tolist()
                    del_rows = recent_out[recent_out["id"].isin(del_ids)]
                    for _, row in del_rows.iterrows():
                        if row["유형"] == "출고":
                            restore_fifo(int(row["id"]))
                        execute("DELETE FROM stock_out WHERE id=?", (int(row["id"]),))
                    st.success(f"{len(to_delete)}건 삭제 완료 (재고 복원됨)")
                    st.rerun()
                else:
                    st.info("삭제할 항목의 '삭제' 체크박스를 선택해주세요.")

    # ── 출고 수정 ──
    with tab3:
        st.markdown('<div class="card">', unsafe_allow_html=True)
        st.markdown("**출고 내역 수기 수정**")

        all_names = ["전체"] + (clients["name"].tolist() if not clients.empty else [])
        fc1, fc2 = st.columns([2, 3])
        sel_client = fc1.selectbox("거래처 필터", all_names, key="edit_filter_client")

        df_edit = run_sql("""
            SELECT o.id, o.date as 날짜, c.name as 거래처, p.name as 품목,
                   o.type as 유형, o.quantity as 수량,
                   ROUND((o.discount_rate*100)::NUMERIC,1) as 할인율PCT,
                   o.fifo_purchase_price as 입고단가,
                   COALESCE(o.fifo_purchase_price,0) * o.quantity as 공급가액,
                   o.total_amount as 납품금액,
                   o.margin_amount as 마진액,
                   o.margin_rate as 마진율, COALESCE(o.note,'') as 메모
            FROM stock_out o
            JOIN clients c ON o.client_id=c.id
            JOIN products p ON o.product_id=p.id
            ORDER BY o.date DESC, o.id
        """)
        if sel_client != "전체":
            df_edit = df_edit[df_edit["거래처"] == sel_client].reset_index(drop=True)

        if df_edit.empty:
            st.info("출고 데이터가 없습니다.")
        else:
            total_sales = int(df_edit["납품금액"].fillna(0).sum())
            st.caption(f"{len(df_edit)}건 | 납품금액 합계: ₩{total_sales:,}")

            edited_df = st.data_editor(
                df_edit,
                disabled=["id", "날짜", "거래처", "품목", "유형", "수량", "할인율PCT", "공급가액"],
                column_config={
                    "id":       st.column_config.NumberColumn("ID", width="small"),
                    "할인율PCT": st.column_config.NumberColumn("할인율%", format="%.1f%%", width="small"),
                    "입고단가": st.column_config.NumberColumn("입고단가(VAT포함) ✏️", format="%d"),
                    "공급가액": st.column_config.NumberColumn("공급가액(VAT포함)", format="%d",
                                                             help="입고단가 × 수량 (자동계산)"),
                    "납품금액": st.column_config.NumberColumn("납품금액(VAT포함) ✏️", format="%d"),
                    "마진액":   st.column_config.NumberColumn("마진액 ✏️", format="%d"),
                    "마진율":   st.column_config.NumberColumn("마진율% ✏️", format="%.2f"),
                },
                use_container_width=True,
                hide_index=True,
                num_rows="fixed",
                key="edit_out_table",
            )

            btn1, btn2 = st.columns(2)
            if btn1.button("💾 변경사항 저장", type="primary", use_container_width=True, key="save_out_edit"):
                conn = get_conn()
                for i, row in edited_df.iterrows():
                    orig = df_edit.iloc[i]
                    changed = (
                        int(row["입고단가"]) != int(orig["입고단가"]) or
                        int(row["납품금액"]) != int(orig["납품금액"]) or
                        int(row["마진액"])   != int(orig["마진액"])   or
                        abs(float(row["마진율"]) - float(orig["마진율"])) > 0.001 or
                        str(row["메모"])     != str(orig["메모"])
                    )
                    if changed:
                        납품금액 = int(row["납품금액"])
                        conn.execute("""
                            UPDATE stock_out SET
                                fifo_purchase_price=?, supply_amount=?, vat_amount=?,
                                total_amount=?, margin_amount=?, margin_rate=?, note=?
                            WHERE id=?
                        """, (int(row["입고단가"]), 납품금액, 0,
                              납품금액, int(row["마진액"]), float(row["마진율"]),
                              str(row["메모"]), int(row["id"])))
                conn.commit(); conn.close()
                st.success("저장 완료!"); st.rerun()

            if btn2.button("🔄 납품가 기준 전체 재계산", use_container_width=True, key="recalc_all_out",
                           help="납품가(VAT포함) × (1-할인율) → 공급가액(VAT포함)·마진 재계산"):
                conn = get_conn()
                records = conn.execute("""
                    SELECT o.id, o.quantity, o.discount_rate, o.fifo_purchase_price,
                           COALESCE(NULLIF(p.settlement_price,0), p.sale_price) as base_p
                    FROM stock_out o JOIN products p ON o.product_id=p.id
                    WHERE o.type='출고'
                """).fetchall()
                for record in records:
                    oid, qty, disc, fifo, sale_p = (record["id"], record["quantity"],
                        record["discount_rate"], record["fifo_purchase_price"], record["base_p"])
                    if sale_p <= 0:
                        continue
                    new_납품금액 = int(sale_p * qty * (1 - disc))   # 납품금액(VAT포함, 판매가 합계)
                    new_공급가액 = fifo * qty                         # 공급가액(VAT포함, 매입원가 합계)
                    new_margin   = new_납품금액 - new_공급가액
                    new_rate     = (new_margin / new_납품금액 * 100) if new_납품금액 > 0 else 0
                    conn.execute("""
                        UPDATE stock_out SET supply_amount=?, vat_amount=?, total_amount=?,
                            margin_amount=?, margin_rate=? WHERE id=?
                    """, (new_납품금액, 0, new_납품금액, new_margin, new_rate, oid))
                conn.commit(); conn.close()
                st.success("전체 재계산 완료!"); st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════
# 미수금 현황
# ══════════════════════════════════════════════════════════════════════════
def page_receivables():

    # 공통 데이터 로드
    recv = run_sql("""
        SELECT c.id, c.name as 거래처,
               COALESCE(s.total_sales,0) as 총매출,
               COALESCE(p.total_paid,0) as 총수금
        FROM clients c
        LEFT JOIN (
            SELECT client_id, SUM(total_amount) as total_sales
            FROM stock_out WHERE type='출고' GROUP BY client_id
        ) s ON c.id=s.client_id
        LEFT JOIN (
            SELECT client_id, SUM(amount) as total_paid
            FROM payments GROUP BY client_id
        ) p ON c.id=p.client_id
        ORDER BY (COALESCE(s.total_sales,0)-COALESCE(p.total_paid,0)) DESC
    """)
    recv["미수금"] = recv["총매출"] - recv["총수금"]
    total_recv = recv["미수금"].sum()
    st.metric("전체 미수금 합계", fmt(total_recv))

    tab1, tab2, tab3, tab4 = st.tabs(["🏥 의원별 미수금 현황", "💳 수금 등록", "📋 거래처 원장", "🥧 미수금 비중"])

    with tab1:
        display = recv[recv["미수금"] > 0].copy()
        if display.empty:
            st.success("미수금 없음")
        else:
            display["총매출"] = display["총매출"].apply(fmt)
            display["총수금"] = display["총수금"].apply(fmt)
            display["미수금"] = display["미수금"].apply(fmt)
            st.dataframe(display[["거래처","총매출","총수금","미수금"]],
                         use_container_width=True, hide_index=True)

    with tab2:
        clients_pay = run_sql("SELECT id, name FROM clients ORDER BY name")
        if clients_pay.empty:
            st.warning("거래처를 먼저 등록해주세요.")
        else:
            pay_client = st.selectbox("거래처", clients_pay["name"].tolist(), key="pay_sel_client")
            client_id  = int(clients_pay[clients_pay["name"]==pay_client].iloc[0]["id"])

            # 현재 미수금 표시
            row = recv[recv["거래처"] == pay_client]
            outstanding = int(row["미수금"].iloc[0]) if not row.empty else 0
            if outstanding > 0:
                st.info(f"현재 미수금: **{fmt(outstanding)}**")
            else:
                st.success("현재 미수금: 없음")

            st.markdown('<div class="card">', unsafe_allow_html=True)
            with st.form("form_pay"):
                col1, col2 = st.columns(2)
                pay_date   = col1.date_input("수금일", value=date.today())
                pay_amount = col2.number_input("수금액", min_value=0, step=10000)
                pay_note   = st.text_input("메모")
                if st.form_submit_button("✅ 수금 등록", use_container_width=True):
                    execute("INSERT INTO payments (date,client_id,amount,note) VALUES (?,?,?,?)",
                            (str(pay_date), client_id, pay_amount, pay_note))
                    st.success(f"수금 완료 — {pay_client} {fmt(pay_amount)}")
                    st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

    with tab3:
        clients_list = run_sql("SELECT id, name FROM clients")
        if clients_list.empty:
            st.info("거래처가 없습니다.")
        else:
            sel_client = st.selectbox("거래처 선택", clients_list["name"].tolist(), key="recv_sel_client")
            cid = int(clients_list[clients_list["name"]==sel_client].iloc[0]["id"])

            # ── 거래 원장 ──
            detail = run_sql("""
                SELECT date as 날짜, '출고' as 구분, total_amount as 금액, 0 as 수금액
                FROM stock_out WHERE client_id=? AND type='출고'
                UNION ALL
                SELECT date, '수금', 0, amount FROM payments WHERE client_id=?
                ORDER BY 날짜 ASC
            """, (cid, cid))
            if not detail.empty:
                detail["금액"]  = detail["금액"].apply(fmt)
                detail["수금액"] = detail["수금액"].apply(fmt)
                st.dataframe(detail, use_container_width=True, hide_index=True)
            else:
                st.info("거래 내역이 없습니다.")

            # ── 수금 내역 수정 / 삭제 ──
            st.divider()
            st.markdown("**✏️ 수금 내역 수정 / 삭제**")
            pay_df = run_sql("""
                SELECT id, date as 날짜, amount as 수금액, COALESCE(note,'') as 메모
                FROM payments WHERE client_id=? ORDER BY date DESC
            """, (cid,))

            if pay_df.empty:
                st.info("등록된 수금 내역이 없습니다.")
            else:
                edited_pay = st.data_editor(
                    pay_df,
                    use_container_width=True,
                    hide_index=True,
                    num_rows="fixed",
                    column_config={
                        "id":    st.column_config.NumberColumn("ID", disabled=True, width="small"),
                        "날짜":  st.column_config.TextColumn("날짜 ✏️", width="small"),
                        "수금액": st.column_config.NumberColumn("수금액 ✏️", format="%d", width="medium"),
                        "메모":  st.column_config.TextColumn("메모 ✏️"),
                    },
                    key="edit_pay_table",
                )

                pa, pb = st.columns(2)
                if pa.button("💾 수정 저장", type="primary", use_container_width=True, key="save_pay_edit"):
                    conn = get_conn()
                    changed = 0
                    for i, row in edited_pay.iterrows():
                        orig = pay_df.iloc[i]
                        if (str(row["날짜"]) != str(orig["날짜"]) or
                                int(row["수금액"]) != int(orig["수금액"]) or
                                str(row["메모"]) != str(orig["메모"])):
                            conn.execute(
                                "UPDATE payments SET date=?, amount=?, note=? WHERE id=?",
                                (str(row["날짜"]), int(row["수금액"]), str(row["메모"]), int(row["id"]))
                            )
                            changed += 1
                    conn.commit(); conn.close()
                    if changed:
                        st.success(f"{changed}건 수정 완료!")
                        st.rerun()
                    else:
                        st.info("변경된 내용이 없습니다.")

                del_options = pay_df.apply(
                    lambda r: f"{r['날짜']} | {fmt(r['수금액'])} | {r['메모']}", axis=1
                ).tolist()
                del_sels = st.multiselect("삭제할 수금 내역 선택", del_options, key="del_pay_sel")
                if del_sels:
                    if pb.button(f"🗑 {len(del_sels)}건 삭제", type="primary",
                                 use_container_width=True, key="del_pay_btn"):
                        del_ids = [int(pay_df.iloc[del_options.index(s)]["id"]) for s in del_sels]
                        conn = get_conn()
                        for pid in del_ids:
                            conn.execute("DELETE FROM payments WHERE id=?", (pid,))
                        conn.commit(); conn.close()
                        st.success(f"{len(del_ids)}건 삭제 완료!")
                        st.rerun()

    with tab4:
        pie_data = recv[recv["미수금"] > 0]
        if pie_data.empty:
            st.info("미수금 데이터 없음")
        else:
            fig = px.pie(pie_data, values="미수금", names="거래처",
                         title="거래처별 미수금 비중", template="plotly_dark",
                         color_discrete_sequence=px.colors.sequential.Plasma_r)
            fig.update_layout(paper_bgcolor="rgba(0,0,0,0)")
            st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})

# ══════════════════════════════════════════════════════════════════════════
# 기초 데이터 관리
# ══════════════════════════════════════════════════════════════════════════
def page_master():
    tab1, tab2 = st.tabs(["🏷️ 제품 관리", "🏢 거래처 관리"])

    with tab1:
        with st.expander("➕ 제품 등록"):
            with st.form("form_prod"):
                r1c1, r1c2, r1c3 = st.columns([3, 2, 2])
                pname  = r1c1.text_input("제품명 *")
                pgroup = r1c2.selectbox("제품군", ["PHA물톡스", "로즈 드 메르"])
                spec   = r1c3.text_input("규격")
                r2c1, r2c2, r2c3 = st.columns([2, 2, 2])
                purchase = r2c1.number_input("공급가", min_value=0, step=100)
                sale     = r2c2.number_input("납품가", min_value=0, step=100)
                minstk   = r2c3.number_input("최소 재고 알림", min_value=0, value=10)
                if st.form_submit_button("등록", use_container_width=True):
                    if pname:
                        execute(
                            "INSERT INTO products (name,product_group,spec,purchase_price,sale_price,settlement_price,min_stock) VALUES (?,?,?,?,?,?,?)",
                            (pname, pgroup, spec, purchase, sale, sale, minstk)
                        )
                        st.success(f"'{pname}' 등록 완료"); st.rerun()
        prods = run_sql("""
            SELECT p.id, p.name as 제품명, p.product_group as 제품군, p.spec as 규격,
                   p.purchase_price as 공급가,
                   p.settlement_price as 납품가,
                   p.min_stock as 최소재고,
                   COALESCE(SUM(s.remaining_qty),0) as 현재고,
                   p.bundle_with_id,
                   COALESCE(b.name, '') as 세트연동
            FROM products p
            LEFT JOIN stock_in s ON p.id=s.product_id
            LEFT JOIN products b ON p.bundle_with_id=b.id
            GROUP BY p.id, p.name, p.product_group, p.spec,
                     p.purchase_price, p.settlement_price, p.min_stock,
                     p.bundle_with_id, b.name
        """)
        PRODUCT_GROUPS = ["PHA물톡스", "로즈 드 메르"]

        if not prods.empty:
            display_prods = prods.drop(columns=["id", "bundle_with_id"])
            edited_prods = st.data_editor(
                display_prods,
                use_container_width=True, hide_index=True, num_rows="fixed",
                column_order=["제품명", "제품군", "규격", "공급가", "납품가", "최소재고", "현재고", "세트연동"],
                column_config={
                    "제품명":   st.column_config.TextColumn("제품명 ✏️"),
                    "제품군":   st.column_config.SelectboxColumn("제품군 ✏️", options=PRODUCT_GROUPS,
                                                                help="PHA물톡스 / 로즈 드 메르"),
                    "규격":     st.column_config.TextColumn("규격 ✏️"),
                    "공급가":   st.column_config.NumberColumn("공급가 ✏️", format="₩%d",
                                                             min_value=0, step=100,
                                                             help="VAT 제외 원가(FIFO 기준 마진 계산용)"),
                    "납품가":   st.column_config.NumberColumn("납품가 ✏️", format="₩%d",
                                                             min_value=0, step=100,
                                                             help="출고 시 할인율 계산 기준가"),
                    "최소재고":  st.column_config.NumberColumn("최소재고", min_value=0, step=1),
                    "현재고":   st.column_config.NumberColumn("현재고", disabled=True),
                    "세트연동":  st.column_config.TextColumn("세트연동", disabled=True,
                                                             help="이 제품 출고 시 함께 자동 출고되는 세트 구성품"),
                },
            )
            if st.button("저장", type="primary", key="save_prods"):
                conn = get_conn()
                for i, row in edited_prods.iterrows():
                    pid     = int(prods.iloc[i]["id"])
                    v_name  = str(row["제품명"]).strip()  if pd.notna(row["제품명"])  else prods.iloc[i]["제품명"]
                    v_group = str(row["제품군"]).strip()  if pd.notna(row["제품군"])  else ""
                    v_spec  = str(row["규격"]).strip()    if pd.notna(row["규격"])    else ""
                    v_purchase = int(row["공급가"])  if pd.notna(row["공급가"])  else 0
                    v_sale     = int(row["납품가"])  if pd.notna(row["납품가"])  else 0
                    v_min      = int(row["최소재고"]) if pd.notna(row["최소재고"]) else 0
                    conn.execute(
                        "UPDATE products SET name=?, product_group=?, spec=?, purchase_price=?, sale_price=?, settlement_price=?, min_stock=? WHERE id=?",
                        (v_name, v_group, v_spec, v_purchase, v_sale, v_sale, v_min, pid)
                    )
                conn.commit(); conn.close()
                st.success("저장 완료!"); st.rerun()

            st.divider()
            with st.expander("🗑️ 제품 삭제"):
                del_sels = st.multiselect("삭제할 제품 선택 (복수 가능)", prods["제품명"].tolist(), key="del_prod_sel")
                if del_sels:
                    conn_chk = get_conn()
                    warn_items, safe_items = [], []
                    for name in del_sels:
                        pid = int(prods[prods["제품명"] == name].iloc[0]["id"])
                        in_cnt  = conn_chk.execute("SELECT COUNT(*) as cnt FROM stock_in  WHERE product_id=?", (pid,)).fetchone()["cnt"]
                        out_cnt = conn_chk.execute("SELECT COUNT(*) as cnt FROM stock_out WHERE product_id=?", (pid,)).fetchone()["cnt"]
                        if in_cnt > 0 or out_cnt > 0:
                            warn_items.append(f"• {name} (입고 {in_cnt}건 / 출고 {out_cnt}건)")
                        else:
                            safe_items.append(f"• {name}")
                    conn_chk.close()
                    if safe_items:
                        st.info("입출고 데이터 없음 (안전):\n" + "\n".join(safe_items))
                    if warn_items:
                        st.warning("관련 입출고 내역도 함께 삭제됩니다:\n" + "\n".join(warn_items))
                    if st.button(f"🗑️ {len(del_sels)}개 삭제 확인", type="primary", key="del_prod_btn"):
                        conn_del = get_conn()
                        for name in del_sels:
                            pid = int(prods[prods["제품명"] == name].iloc[0]["id"])
                            conn_del.execute("DELETE FROM stock_in  WHERE product_id=?", (pid,))
                            conn_del.execute("DELETE FROM stock_out WHERE product_id=?", (pid,))
                            conn_del.execute("UPDATE products SET bundle_with_id=NULL WHERE bundle_with_id=?", (pid,))
                            conn_del.execute("DELETE FROM products WHERE id=?", (pid,))
                        conn_del.commit(); conn_del.close()
                        st.success(f"{len(del_sels)}개 제품 삭제 완료"); st.rerun()

            st.divider()
            with st.expander("🔗 세트 상품 설정"):
                st.caption("출고 등록 시 주 제품을 추가하면 연동 제품이 자동으로 함께 추가됩니다.")
                all_prod_names = ["(없음)"] + prods["제품명"].tolist()
                bc1, bc2, bc3 = st.columns([2, 2, 1])
                bundle_main = bc1.selectbox("주 제품", prods["제품명"].tolist(), key="bundle_main_sel")
                bundle_comp = bc2.selectbox("연동 제품 (세트 구성품)", all_prod_names, key="bundle_comp_sel")
                if bc3.button("설정 저장", type="primary", use_container_width=True, key="bundle_save"):
                    main_id = int(prods[prods["제품명"] == bundle_main].iloc[0]["id"])
                    if bundle_comp == "(없음)":
                        execute("UPDATE products SET bundle_with_id=NULL WHERE id=?", (main_id,))
                        st.success(f"'{bundle_main}' 세트 설정 해제")
                    else:
                        comp_id = int(prods[prods["제품명"] == bundle_comp].iloc[0]["id"])
                        if main_id == comp_id:
                            st.error("주 제품과 연동 제품이 동일합니다.")
                        else:
                            execute("UPDATE products SET bundle_with_id=? WHERE id=?", (comp_id, main_id))
                            st.success(f"'{bundle_main}' → '{bundle_comp}' 세트 설정 완료")
                    st.rerun()
                current_bundles = prods[prods["세트연동"] != ""][["제품명", "세트연동"]]
                if not current_bundles.empty:
                    st.markdown("**현재 세트 설정:**")
                    for _, brow in current_bundles.iterrows():
                        st.caption(f"• {brow['제품명']}  →  {brow['세트연동']}")

    with tab2:
        with st.expander("➕ 거래처 등록"):
            with st.form("form_client"):
                r1c1, r1c2, r1c3 = st.columns([2, 1, 1])
                cname    = r1c1.text_input("거래처명 *")
                category = r1c2.selectbox("대분류", CATEGORIES,
                                          help="선택하면 기본 할인율 자동 적용")
                contact  = r1c3.text_input("담당자")
                base_disc = int(CATEGORY_DISCOUNT.get(category, 0.17) * 100) if category else 17
                if category:
                    st.info(f"[{category}] 기본 할인율: {int(CATEGORY_DISCOUNT[category]*100)}%")
                disc = st.slider("세부 할인율 조정 (%)", 0, 50, base_disc,
                                 help="대분류 기본값에서 개별 조정 가능")
                if st.form_submit_button("등록", use_container_width=True):
                    if cname:
                        execute("""INSERT INTO clients
                                   (name,category,contact,discount_rate)
                                   VALUES (?,?,?,?)""",
                                (cname, category, contact, disc/100))
                        st.success(f"거래처 '{cname}' 등록 완료 (할인율 {disc}%)")
                        st.rerun()
        st.divider()
        st.subheader("등록된 거래처 수정")
        clients_df = run_sql("""
            SELECT id, name as 거래처명, category as 대분류,
                   contact as 담당자,
                   ROUND((discount_rate*100)::NUMERIC,1) as 할인율_PCT
            FROM clients ORDER BY name
        """)
        if not clients_df.empty:
            edited = st.data_editor(
                clients_df,
                use_container_width=True, hide_index=True, num_rows="fixed",
                column_config={
                    "id":         st.column_config.NumberColumn("ID", disabled=True, width="small"),
                    "거래처명":   st.column_config.TextColumn("거래처명", width="medium"),
                    "대분류":     st.column_config.SelectboxColumn("대분류", options=CATEGORIES, width="small"),
                    "담당자":     st.column_config.TextColumn("담당자", width="small"),
                    "할인율_PCT": st.column_config.NumberColumn("할인율(%)", min_value=0, max_value=50,
                                                                step=0.5, width="small",
                                                                help="숫자 직접 입력 또는 대분류 변경 후 저장"),
                },
            )
            st.caption("* 대분류를 변경해도 할인율(%)은 자동 적용되지 않습니다. 할인율(%)을 직접 수정하거나 아래 버튼을 눌러주세요.")
            col_a, col_b = st.columns([1, 3])
            with col_a:
                if st.button("대분류 기본 할인율 일괄 적용", use_container_width=True):
                    for _, row in edited.iterrows():
                        cat = row["대분류"]
                        if cat and cat in CATEGORY_DISCOUNT:
                            edited.loc[edited["id"]==row["id"], "할인율_PCT"] = CATEGORY_DISCOUNT[cat]*100
            if st.button("변경사항 저장", type="primary", use_container_width=True):
                conn = get_conn()
                for _, row in edited.iterrows():
                    conn.execute("""UPDATE clients
                                    SET name=?, category=?, contact=?, discount_rate=?
                                    WHERE id=?""",
                                 (row["거래처명"], row["대분류"] or "",
                                  row["담당자"] or "", float(row["할인율_PCT"])/100,
                                  int(row["id"])))
                conn.commit(); conn.close()
                st.success("거래처 정보 저장 완료!"); st.rerun()

# ══════════════════════════════════════════════════════════════════════════
# 의료장비 — 리포트
# ══════════════════════════════════════════════════════════════════════════
def page_med_report():
    today       = date.today()
    today_str   = today.strftime("%Y-%m-%d")
    month_start = today.replace(day=1).strftime("%Y-%m-%d")
    year_start  = today.replace(month=1, day=1).strftime("%Y-%m-%d")
    this_month  = today.strftime("%Y-%m")

    kpi = run_sql("""
        SELECT
            COALESCE(SUM(CASE WHEN date=? THEN total_amount END),0)              as today,
            COALESCE(SUM(CASE WHEN date BETWEEN ? AND ? THEN total_amount END),0) as month,
            COALESCE(SUM(CASE WHEN date BETWEEN ? AND ? THEN total_amount END),0) as year,
            COUNT(CASE WHEN date BETWEEN ? AND ? THEN 1 END)                      as month_cnt
        FROM med_stock_out
    """, (today_str, month_start, today_str, year_start, today_str, month_start, today_str)).iloc[0]

    total_out  = run_sql("SELECT COALESCE(SUM(total_amount),0) as v FROM med_stock_out").iloc[0]["v"]
    total_paid = run_sql("SELECT COALESCE(SUM(amount),0) as v FROM med_payments").iloc[0]["v"]
    total_recv = total_out - total_paid

    st.markdown(f"""
    <div class="kpi-row">
      <div class="kpi"><div class="kpi-label">오늘 납품액</div><div class="kpi-value">{fmt(kpi['today'])}</div><div class="kpi-sub">{today_str}</div></div>
      <div class="kpi"><div class="kpi-label">이번달 납품액</div><div class="kpi-value">{fmt(kpi['month'])}</div><div class="kpi-sub">{this_month}</div></div>
      <div class="kpi"><div class="kpi-label">연간 납품액</div><div class="kpi-value">{fmt(kpi['year'])}</div><div class="kpi-sub">{year_start[:4]}년</div></div>
      <div class="kpi"><div class="kpi-label">전체 미수금</div><div class="kpi-value">{fmt(total_recv)}</div><div class="kpi-sub">누적</div></div>
    </div>
    """, unsafe_allow_html=True)

    gc1, gc2 = st.columns(2)
    with gc1:
        monthly = run_sql("SELECT substr(date,1,7) as 월, SUM(total_amount) as 납품액 FROM med_stock_out GROUP BY 월 ORDER BY 월")
        if not monthly.empty:
            fig = px.bar(monthly, x="월", y="납품액", title="월별 납품액 추이", template="plotly_dark",
                         color_discrete_sequence=["#0d9488"])
            fig.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                              margin=dict(l=0, r=0, t=30, b=0))
            st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
        else:
            st.info("납품 데이터가 없습니다.")
    with gc2:
        pm = run_sql("SELECT payment_method as 결제방식, SUM(total_amount) as 금액 FROM med_stock_out GROUP BY payment_method")
        if not pm.empty:
            fig2 = px.pie(pm, values="금액", names="결제방식", title="결제방식별 비중",
                          template="plotly_dark", color_discrete_sequence=px.colors.sequential.Teal)
            fig2.update_layout(paper_bgcolor="rgba(0,0,0,0)", margin=dict(l=0, r=0, t=30, b=0))
            st.plotly_chart(fig2, use_container_width=True, config={"displayModeBar": False})

    st.subheader("병원별 납품 현황")
    hosp = run_sql("""
        SELECT c.hospital_name as 병원명,
               COUNT(o.id) as 납품건수,
               COALESCE(SUM(o.total_amount),0) as 납품액합계,
               COALESCE(SUM(p.amount),0) as 수금액,
               COALESCE(SUM(o.total_amount),0)-COALESCE(SUM(p.amount),0) as 미수금
        FROM med_clients c
        LEFT JOIN med_stock_out o ON c.id=o.client_id
        LEFT JOIN med_payments  p ON c.id=p.client_id
        GROUP BY c.id, c.hospital_name ORDER BY 납품액합계 DESC
    """)
    if not hosp.empty:
        d = hosp.copy()
        for col in ["납품액합계","수금액","미수금"]:
            d[col] = d[col].apply(fmt)
        st.dataframe(d, use_container_width=True, hide_index=True)

    st.markdown('<div class="card"><div class="card-title">📥 엑셀 다운로드</div>', unsafe_allow_html=True)
    xc1, xc2, xc3 = st.columns([2, 2, 2])
    xl_from = xc1.date_input("시작일", value=date(today.year, 1, 1), key="med_xl_from")
    xl_to   = xc2.date_input("종료일", value=today, key="med_xl_to")
    if xc3.button("📊 엑셀 생성", type="primary", use_container_width=True, key="med_gen_xl"):
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        frm, to_ = str(xl_from), str(xl_to)
        wb = Workbook()
        hf = Font(bold=True, color="FFFFFF", size=10)
        hfl = PatternFill("solid", fgColor="0F766E")
        th = Side(style="thin", color="CCCCCC")
        bd = Border(left=th, right=th, top=th, bottom=th)
        ctr = Alignment(horizontal="center", vertical="center")
        def sh(ws, row, cols):
            for c in range(1, cols+1):
                cell = ws.cell(row=row, column=c)
                cell.font=hf; cell.fill=hfl; cell.alignment=ctr; cell.border=bd
        def aw(ws):
            for col in ws.columns:
                ml = max((len(str(cell.value or "")) for cell in col), default=0)
                ws.column_dimensions[get_column_letter(col[0].column)].width = min(ml+4, 30)

        df_out = run_sql("""
            SELECT o.date as 날짜, c.hospital_name as 병원명, o.product_name as 장비명,
                   o.manufacturer as 장비사, o.serial_number as 시리얼번호,
                   o.is_purchased as 매입여부, o.payment_method as 결제방식,
                   o.sale_price as 납품가, o.commission as 수수료, o.total_amount as 합계
            FROM med_stock_out o LEFT JOIN med_clients c ON o.client_id=c.id
            WHERE o.date BETWEEN ? AND ? ORDER BY o.date
        """, (frm, to_))
        ws1 = wb.active; ws1.title = "납품 내역"
        cols1 = ["날짜","병원명","장비명","장비사","시리얼번호","매입여부","결제방식","납품가","수수료","합계"]
        ws1.append(cols1); sh(ws1, 1, len(cols1))
        for _, r in df_out.iterrows():
            ws1.append([r["날짜"],r["병원명"],r["장비명"],r["장비사"],r["시리얼번호"],
                        r["매입여부"],r["결제방식"],int(r["납품가"]),int(r["수수료"]),int(r["합계"])])
        aw(ws1)

        df_hosp = run_sql("""
            SELECT c.hospital_name as 병원명, COUNT(o.id) as 납품건수,
                   COALESCE(SUM(o.total_amount),0) as 납품액합계,
                   COALESCE(SUM(p.amount),0) as 수금액,
                   COALESCE(SUM(o.total_amount),0)-COALESCE(SUM(p.amount),0) as 미수금
            FROM med_clients c
            LEFT JOIN med_stock_out o ON c.id=o.client_id AND o.date BETWEEN ? AND ?
            LEFT JOIN med_payments  p ON c.id=p.client_id AND p.date BETWEEN ? AND ?
            GROUP BY c.id, c.hospital_name ORDER BY 납품액합계 DESC
        """, (frm, to_, frm, to_))
        ws2 = wb.create_sheet("병원별 요약")
        cols2 = ["병원명","납품건수","납품액합계","수금액","미수금"]
        ws2.append(cols2); sh(ws2, 1, len(cols2))
        for _, r in df_hosp.iterrows():
            ws2.append([r["병원명"],int(r["납품건수"]),int(r["납품액합계"]),int(r["수금액"]),int(r["미수금"])])
        aw(ws2)

        buf2 = io.BytesIO(); wb.save(buf2); buf2.seek(0)
        st.download_button("⬇ 엑셀 다운로드", data=buf2,
                           file_name=f"의료장비_리포트_{frm}_{to_}.xlsx",
                           mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                           use_container_width=True)
    st.markdown("</div>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════
# 의료장비 — 영업 관리
# ══════════════════════════════════════════════════════════════════════════
def page_med_sales():
    tab1, tab2 = st.tabs(["📥 매입 등록", "📤 납품 등록"])

    # ── 매입 등록 ──
    with tab1:
        ep = run_sql("SELECT id, name, manufacturer FROM med_products ORDER BY name")
        mfrs = sorted(ep["manufacturer"].dropna().unique().tolist()) if not ep.empty else []

        st.markdown('<div class="card">', unsafe_allow_html=True)
        in_mode = st.radio("입력 방식", ["기존 선택", "직접 입력"], horizontal=True, key="med_in_mode")
        in_date = st.date_input("매입일", value=date.today(), key="med_in_date")

        if in_mode == "기존 선택" and not ep.empty:
            c1, c2 = st.columns(2)
            sel_mfr = c1.selectbox("장비사", ["전체"] + mfrs, key="med_in_mfr_sel")
            filtered = ep if sel_mfr == "전체" else ep[ep["manufacturer"] == sel_mfr]
            sel_prod = c2.selectbox("장비명", filtered["name"].tolist() if not filtered.empty else [], key="med_in_prod_sel")
            if not filtered.empty and sel_prod:
                row_p = filtered[filtered["name"] == sel_prod]
                manufacturer = row_p["manufacturer"].iloc[0] if not row_p.empty else ""
            else:
                manufacturer = ""
            product_name = sel_prod
        else:
            c1, c2 = st.columns(2)
            manufacturer  = c1.text_input("장비사", key="med_in_mfr_txt")
            product_name  = c2.text_input("장비명", key="med_in_prod_txt")

        c3, c4 = st.columns(2)
        price_type     = c3.radio("가격 유형", ["직납", "매입"], horizontal=True, key="med_in_pt")
        purchase_price = c4.number_input("장비 가격", min_value=0, step=100000, key="med_in_price")

        c5, c6 = st.columns(2)
        serial      = c5.text_input("시리얼번호", key="med_in_serial")
        warranty    = c6.date_input("보증만료일", value=None, key="med_in_warranty")
        note_in     = st.text_input("비고", key="med_in_note")

        if st.button("✅ 매입 등록", type="primary", use_container_width=True, key="med_in_submit"):
            if not manufacturer or not product_name:
                st.error("장비사와 장비명을 입력해주세요.")
            else:
                exists = run_sql("SELECT id FROM med_products WHERE name=? AND manufacturer=?", (product_name, manufacturer))
                if exists.empty:
                    execute("INSERT INTO med_products (name, manufacturer) VALUES (?,?)", (product_name, manufacturer))
                execute("""INSERT INTO med_stock_in
                    (date,manufacturer,product_name,price_type,purchase_price,serial_number,warranty_end,note)
                    VALUES (?,?,?,?,?,?,?,?)""",
                    (str(in_date), manufacturer, product_name, price_type, purchase_price,
                     serial, str(warranty) if warranty else None, note_in))
                st.success(f"매입 등록 완료 — {manufacturer} {product_name}"); st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

        df_in = run_sql("""SELECT id, date as 날짜, manufacturer as 장비사, product_name as 장비명,
                                  price_type as 가격유형, purchase_price as 매입가,
                                  COALESCE(serial_number,'') as 시리얼번호,
                                  COALESCE(warranty_end,'') as 보증만료일, COALESCE(note,'') as 비고
                           FROM med_stock_in ORDER BY date DESC""")
        if not df_in.empty:
            st.subheader("전체 매입 내역")
            disp = df_in.copy(); disp.insert(0, "삭제", False)
            edited_in = st.data_editor(disp, use_container_width=True, hide_index=True, num_rows="fixed",
                disabled=["날짜","장비사","장비명","가격유형","매입가","시리얼번호","보증만료일","비고"],
                column_config={"id": None, "삭제": st.column_config.CheckboxColumn("삭제", width="small"),
                               "매입가": st.column_config.NumberColumn("매입가", format="₩%d")},
                key="med_in_hist")
            if st.button("🗑 삭제 적용", key="med_in_del", type="primary"):
                to_del = edited_in[edited_in["삭제"] == True]
                if not to_del.empty:
                    conn = get_conn()
                    for did in to_del["id"].astype(int):
                        conn.execute("DELETE FROM med_stock_in WHERE id=?", (did,))
                    conn.commit(); conn.close()
                    st.success(f"{len(to_del)}건 삭제 완료"); st.rerun()
                else:
                    st.info("삭제할 항목의 체크박스를 선택해주세요.")

    # ── 납품 등록 ──
    with tab2:
        med_clients = run_sql("SELECT id, hospital_name FROM med_clients ORDER BY hospital_name")

        st.markdown('<div class="card">', unsafe_allow_html=True)
        out_date = st.date_input("납품일", value=date.today(), key="med_out_date")

        hosp_mode = st.radio("병원 입력 방식", ["기존 선택", "직접 입력"], horizontal=True, key="med_out_hosp_mode")
        if hosp_mode == "기존 선택" and not med_clients.empty:
            hospital_name = st.selectbox("병원명", med_clients["hospital_name"].tolist(), key="med_out_hosp_sel")
            client_id     = int(med_clients[med_clients["hospital_name"] == hospital_name].iloc[0]["id"])
        else:
            hospital_name = st.text_input("병원명 (직접 입력)", key="med_out_hosp_txt")
            client_id     = None

        oc1, oc2 = st.columns(2)
        product_name_out = oc1.text_input("장비명", key="med_out_prod")
        manufacturer_out = oc2.text_input("장비사", key="med_out_mfr")

        mc1, mc2 = st.columns(2)
        sale_price = mc1.number_input("납품가 (VAT포함)", min_value=0, step=100000, key="med_out_sale")
        commission = mc2.number_input("장비 수수료", min_value=0, step=10000, key="med_out_comm")

        supply_amount = round(sale_price / 1.1) if sale_price > 0 else 0
        vat_amount    = sale_price - supply_amount
        if sale_price > 0:
            st.caption(f"공급가액: {fmt(supply_amount)} | 부가세: {fmt(vat_amount)} | 합계: {fmt(sale_price)}")

        fc1, fc2, fc3 = st.columns(3)
        serial_out     = fc1.text_input("시리얼번호", key="med_out_serial")
        is_purchased   = fc2.radio("매입여부", ["매입", "위탁"], horizontal=True, key="med_out_is_purch")
        payment_method = fc3.radio("결제방식", ["현금", "카드", "리스"], horizontal=True, key="med_out_pay")
        note_out       = st.text_input("비고", key="med_out_note")

        st.markdown("**계약서 파일 첨부** (선택)")
        contract_file = st.file_uploader("JPG / PNG / PDF", type=["jpg","jpeg","png","pdf"], key="med_out_file")

        if st.button("✅ 납품 등록", type="primary", use_container_width=True, key="med_out_submit"):
            if not hospital_name or not product_name_out:
                st.error("병원명과 장비명을 입력해주세요.")
            else:
                if client_id is None:
                    exc = run_sql("SELECT id FROM med_clients WHERE hospital_name=?", (hospital_name,))
                    if exc.empty:
                        execute("INSERT INTO med_clients (hospital_name) VALUES (?)", (hospital_name,))
                    client_id = int(run_sql("SELECT id FROM med_clients WHERE hospital_name=?", (hospital_name,)).iloc[0]["id"])

                if product_name_out and manufacturer_out:
                    ep2 = run_sql("SELECT id FROM med_products WHERE name=? AND manufacturer=?", (product_name_out, manufacturer_out))
                    if ep2.empty:
                        execute("INSERT INTO med_products (name, manufacturer) VALUES (?,?)", (product_name_out, manufacturer_out))
                prod_row = run_sql("SELECT id FROM med_products WHERE name=?", (product_name_out,))
                prod_id  = int(prod_row.iloc[0]["id"]) if not prod_row.empty else None

                drive_file_id, drive_file_name = None, None
                if contract_file is not None:
                    svc = _gdrive_service()
                    if svc:
                        fname = f"{hospital_name}_{product_name_out}_{out_date}{Path(contract_file.name).suffix}"
                        drive_file_id, _ = drive_upload(contract_file.read(), fname, contract_file.type)
                        drive_file_name  = fname
                    else:
                        st.warning("Google Drive 미연동 — 파일 첨부 건너뜀")

                execute("""INSERT INTO med_stock_out
                    (date,client_id,product_id,product_name,manufacturer,serial_number,
                     is_purchased,payment_method,sale_price,commission,
                     supply_amount,vat_amount,total_amount,drive_file_id,drive_file_name,note)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (str(out_date), client_id, prod_id, product_name_out, manufacturer_out,
                     serial_out, is_purchased, payment_method, sale_price, commission,
                     supply_amount, vat_amount, sale_price, drive_file_id, drive_file_name, note_out))
                st.success(f"납품 등록 완료 — {hospital_name} | {product_name_out}"); st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

        df_out = run_sql("""SELECT o.id, o.date as 날짜, c.hospital_name as 병원명,
                                   o.product_name as 장비명, o.manufacturer as 장비사,
                                   o.is_purchased as 매입여부, o.payment_method as 결제방식,
                                   o.sale_price as 납품가, o.commission as 수수료, o.total_amount as 합계,
                                   COALESCE(o.drive_file_name,'') as 첨부파일, o.drive_file_id
                            FROM med_stock_out o LEFT JOIN med_clients c ON o.client_id=c.id
                            ORDER BY o.date DESC""")
        if not df_out.empty:
            st.subheader("전체 납품 내역")
            disp_out = df_out.drop(columns=["drive_file_id"]).copy()
            disp_out.insert(0, "삭제", False)
            edited_out = st.data_editor(disp_out, use_container_width=True, hide_index=True, num_rows="fixed",
                disabled=["날짜","병원명","장비명","장비사","매입여부","결제방식","납품가","수수료","합계","첨부파일"],
                column_config={"id": None, "삭제": st.column_config.CheckboxColumn("삭제", width="small"),
                               "납품가": st.column_config.NumberColumn("납품가", format="₩%d"),
                               "합계":   st.column_config.NumberColumn("합계",   format="₩%d")},
                key="med_out_hist")
            if st.button("🗑 삭제 적용", key="med_out_del", type="primary"):
                to_del = edited_out[edited_out["삭제"] == True]
                if not to_del.empty:
                    conn = get_conn()
                    for idx in to_del.index:
                        did = int(df_out.loc[idx, "id"])
                        fid = df_out.loc[idx, "drive_file_id"]
                        if fid: drive_delete(fid)
                        conn.execute("DELETE FROM med_stock_out WHERE id=?", (did,))
                    conn.commit(); conn.close()
                    st.success(f"{len(to_del)}건 삭제 완료"); st.rerun()
                else:
                    st.info("삭제할 항목의 체크박스를 선택해주세요.")


# ══════════════════════════════════════════════════════════════════════════
# 의료장비 — 수금 현황
# ══════════════════════════════════════════════════════════════════════════
def page_med_receivables():
    recv = run_sql("""
        SELECT c.id, c.hospital_name as 병원명,
               COALESCE(o.tot,0) as 총납품액,
               COALESCE(p.tot,0) as 총수금
        FROM med_clients c
        LEFT JOIN (SELECT client_id, SUM(total_amount) as tot FROM med_stock_out GROUP BY client_id) o ON c.id=o.client_id
        LEFT JOIN (SELECT client_id, SUM(amount)       as tot FROM med_payments  GROUP BY client_id) p ON c.id=p.client_id
        ORDER BY (COALESCE(o.tot,0)-COALESCE(p.tot,0)) DESC
    """)
    if recv.empty:
        st.info("거래처 데이터가 없습니다. 납품을 먼저 등록해주세요.")
        return
    recv["미수금"] = recv["총납품액"] - recv["총수금"]
    st.metric("전체 미수금 합계", fmt(recv["미수금"].sum()))

    tab1, tab2, tab3 = st.tabs(["🏥 병원별 미수금", "💳 수금 등록", "📋 거래처 원장"])

    with tab1:
        disp = recv[recv["미수금"] > 0].copy()
        if disp.empty:
            st.success("미수금 없음")
        else:
            for col in ["총납품액","총수금","미수금"]:
                disp[col] = disp[col].apply(fmt)
            st.dataframe(disp[["병원명","총납품액","총수금","미수금"]], use_container_width=True, hide_index=True)

        st.divider()
        st.markdown("**리스 납품 현황**")
        lease = run_sql("""SELECT o.date as 날짜, c.hospital_name as 병원명, o.product_name as 장비명,
                                  o.total_amount as 납품가, COALESCE(o.note,'') as 비고
                           FROM med_stock_out o LEFT JOIN med_clients c ON o.client_id=c.id
                           WHERE o.payment_method='리스' ORDER BY o.date DESC""")
        if not lease.empty:
            lease["납품가"] = lease["납품가"].apply(fmt)
            st.dataframe(lease, use_container_width=True, hide_index=True)
        else:
            st.info("리스 납품 건 없음")

    with tab2:
        clients_pay = run_sql("SELECT id, hospital_name FROM med_clients ORDER BY hospital_name")
        if clients_pay.empty:
            st.warning("거래처를 먼저 등록해주세요.")
        else:
            pay_client = st.selectbox("병원 선택", clients_pay["hospital_name"].tolist(), key="med_pay_client")
            cid = int(clients_pay[clients_pay["hospital_name"] == pay_client].iloc[0]["id"])
            row = recv[recv["병원명"] == pay_client]
            outstanding = int(row["미수금"].iloc[0]) if not row.empty else 0
            if outstanding > 0:
                st.info(f"현재 미수금: **{fmt(outstanding)}**")
            else:
                st.success("현재 미수금: 없음")

            st.markdown('<div class="card">', unsafe_allow_html=True)
            with st.form("med_pay_form"):
                pc1, pc2, pc3 = st.columns(3)
                pay_date   = pc1.date_input("수금일", value=date.today())
                pay_amount = pc2.number_input("수금액", min_value=0, step=100000)
                pay_method = pc3.radio("결제방식", ["현금", "카드", "리스"], horizontal=True)
                pay_note   = st.text_input("메모")
                if st.form_submit_button("✅ 수금 등록", use_container_width=True):
                    execute("INSERT INTO med_payments (date,client_id,amount,payment_method,note) VALUES (?,?,?,?,?)",
                            (str(pay_date), cid, pay_amount, pay_method, pay_note))
                    st.success(f"수금 완료 — {pay_client} {fmt(pay_amount)}"); st.rerun()
            st.markdown("</div>", unsafe_allow_html=True)

    with tab3:
        clients_list = run_sql("SELECT id, hospital_name FROM med_clients")
        if clients_list.empty:
            st.info("거래처가 없습니다.")
        else:
            sel = st.selectbox("병원 선택", clients_list["hospital_name"].tolist(), key="med_ledger_sel")
            cid2 = int(clients_list[clients_list["hospital_name"] == sel].iloc[0]["id"])
            detail = run_sql("""
                SELECT date as 날짜, '납품' as 구분, total_amount as 납품액, 0 as 수금액
                FROM med_stock_out WHERE client_id=?
                UNION ALL
                SELECT date, '수금', 0, amount FROM med_payments WHERE client_id=?
                ORDER BY 날짜 ASC
            """, (cid2, cid2))
            if not detail.empty:
                d2 = detail.copy()
                d2["납품액"] = d2["납품액"].apply(fmt)
                d2["수금액"] = d2["수금액"].apply(fmt)
                st.dataframe(d2, use_container_width=True, hide_index=True)
            else:
                st.info("거래 내역 없음")

            st.divider()
            st.markdown("**✏️ 수금 내역 수정/삭제**")
            pay_df = run_sql("""SELECT id, date as 날짜, amount as 수금액,
                                       payment_method as 결제방식, COALESCE(note,'') as 메모
                                FROM med_payments WHERE client_id=? ORDER BY date DESC""", (cid2,))
            if pay_df.empty:
                st.info("등록된 수금 내역이 없습니다.")
            else:
                edited_pay = st.data_editor(pay_df, use_container_width=True, hide_index=True, num_rows="fixed",
                    column_config={
                        "id":     st.column_config.NumberColumn("ID", disabled=True, width="small"),
                        "날짜":   st.column_config.TextColumn("날짜 ✏️"),
                        "수금액": st.column_config.NumberColumn("수금액 ✏️", format="%d"),
                        "결제방식": st.column_config.SelectboxColumn("결제방식", options=["현금","카드","리스"]),
                        "메모":   st.column_config.TextColumn("메모 ✏️"),
                    }, key="med_pay_edit")
                pa, pb = st.columns(2)
                if pa.button("💾 수정 저장", type="primary", use_container_width=True, key="med_pay_save"):
                    conn = get_conn(); changed = 0
                    for i, row in edited_pay.iterrows():
                        orig = pay_df.iloc[i]
                        if (str(row["날짜"]) != str(orig["날짜"]) or
                                int(row["수금액"]) != int(orig["수금액"]) or
                                str(row["메모"]) != str(orig["메모"])):
                            conn.execute("UPDATE med_payments SET date=?,amount=?,payment_method=?,note=? WHERE id=?",
                                        (str(row["날짜"]), int(row["수금액"]), str(row["결제방식"]),
                                         str(row["메모"]), int(row["id"])))
                            changed += 1
                    conn.commit(); conn.close()
                    if changed: st.success(f"{changed}건 수정 완료!"); st.rerun()
                    else: st.info("변경된 내용이 없습니다.")

                del_opts = pay_df.apply(lambda r: f"{r['날짜']} | {fmt(r['수금액'])}", axis=1).tolist()
                del_sels = st.multiselect("삭제할 수금 내역", del_opts, key="med_pay_del_sel")
                if del_sels and pb.button(f"🗑 {len(del_sels)}건 삭제", type="primary",
                                          use_container_width=True, key="med_pay_del_btn"):
                    del_ids = [int(pay_df.iloc[del_opts.index(s)]["id"]) for s in del_sels]
                    conn = get_conn()
                    for pid in del_ids: conn.execute("DELETE FROM med_payments WHERE id=?", (pid,))
                    conn.commit(); conn.close()
                    st.success(f"{len(del_ids)}건 삭제 완료!"); st.rerun()


# ══════════════════════════════════════════════════════════════════════════
# 의료장비 — 기초 관리
# ══════════════════════════════════════════════════════════════════════════
def page_med_master():
    tab1, tab2, tab3, tab4 = st.tabs(["🔧 장비 관리", "🏥 거래처 관리", "📄 계약서 발급", "📁 문서함"])

    # ── 장비 관리 ──
    with tab1:
        with st.expander("➕ 장비 등록"):
            with st.form("med_prod_form"):
                mc1, mc2, mc3 = st.columns(3)
                pname = mc1.text_input("장비명 *")
                mfr   = mc2.text_input("장비사 *")
                spec  = mc3.text_input("규격")
                if st.form_submit_button("등록", use_container_width=True):
                    if pname and mfr:
                        execute("INSERT INTO med_products (name,manufacturer,spec) VALUES (?,?,?)", (pname, mfr, spec))
                        st.success(f"'{pname}' 등록 완료"); st.rerun()
        prods = run_sql("SELECT id, name as 장비명, manufacturer as 장비사, COALESCE(spec,'') as 규격 FROM med_products ORDER BY name")
        if not prods.empty:
            edited_p = st.data_editor(prods, use_container_width=True, hide_index=True, num_rows="fixed",
                column_config={
                    "id":   st.column_config.NumberColumn("ID", disabled=True, width="small"),
                    "장비명": st.column_config.TextColumn("장비명 ✏️"),
                    "장비사": st.column_config.TextColumn("장비사 ✏️"),
                    "규격":   st.column_config.TextColumn("규격 ✏️"),
                }, key="med_prod_edit")
            ca, cb = st.columns(2)
            if ca.button("💾 저장", type="primary", use_container_width=True, key="med_prod_save"):
                conn = get_conn()
                for i, row in edited_p.iterrows():
                    conn.execute("UPDATE med_products SET name=?,manufacturer=?,spec=? WHERE id=?",
                                 (str(row["장비명"]), str(row["장비사"]), str(row["규격"]), int(row["id"])))
                conn.commit(); conn.close(); st.success("저장 완료!"); st.rerun()
            del_p = st.multiselect("삭제할 장비", prods["장비명"].tolist(), key="med_prod_del")
            if del_p and cb.button(f"🗑 {len(del_p)}개 삭제", type="primary", use_container_width=True, key="med_prod_del_btn"):
                conn = get_conn()
                for name in del_p:
                    conn.execute("DELETE FROM med_products WHERE id=?", (int(prods[prods["장비명"]==name].iloc[0]["id"]),))
                conn.commit(); conn.close(); st.success(f"{len(del_p)}개 삭제 완료"); st.rerun()

    # ── 거래처 관리 ──
    with tab2:
        with st.expander("➕ 거래처 등록"):
            with st.form("med_client_form"):
                cc1, cc2 = st.columns(2)
                hname   = cc1.text_input("병원명 *")
                ceo     = cc2.text_input("대표자명")
                cc3, cc4 = st.columns(2)
                contact = cc3.text_input("담당자")
                phone   = cc4.text_input("연락처")
                address = st.text_input("주소")
                credit  = st.number_input("여신한도", min_value=0, step=1000000)
                if st.form_submit_button("등록", use_container_width=True):
                    if hname:
                        execute("INSERT INTO med_clients (hospital_name,ceo_name,contact_person,phone,address,credit_limit) VALUES (?,?,?,?,?,?)",
                                (hname, ceo, contact, phone, address, credit))
                        st.success(f"'{hname}' 등록 완료"); st.rerun()
        clients_df = run_sql("""SELECT id, hospital_name as 병원명, COALESCE(ceo_name,'') as 대표자,
                                       COALESCE(contact_person,'') as 담당자, COALESCE(phone,'') as 연락처,
                                       COALESCE(address,'') as 주소, credit_limit as 여신한도
                                FROM med_clients ORDER BY hospital_name""")
        if not clients_df.empty:
            edited_cl = st.data_editor(clients_df, use_container_width=True, hide_index=True, num_rows="fixed",
                column_config={
                    "id":    st.column_config.NumberColumn("ID", disabled=True, width="small"),
                    "병원명":  st.column_config.TextColumn("병원명 ✏️"),
                    "대표자":  st.column_config.TextColumn("대표자 ✏️"),
                    "담당자":  st.column_config.TextColumn("담당자 ✏️"),
                    "연락처":  st.column_config.TextColumn("연락처 ✏️"),
                    "주소":    st.column_config.TextColumn("주소 ✏️"),
                    "여신한도": st.column_config.NumberColumn("여신한도 ✏️", format="₩%d"),
                }, key="med_client_edit")
            da, db = st.columns(2)
            if da.button("💾 저장", type="primary", use_container_width=True, key="med_client_save"):
                conn = get_conn()
                for i, row in edited_cl.iterrows():
                    conn.execute("""UPDATE med_clients SET hospital_name=?,ceo_name=?,contact_person=?,
                                    phone=?,address=?,credit_limit=? WHERE id=?""",
                                 (str(row["병원명"]), str(row["대표자"]), str(row["담당자"]),
                                  str(row["연락처"]), str(row["주소"]), int(row["여신한도"]), int(row["id"])))
                conn.commit(); conn.close(); st.success("저장 완료!"); st.rerun()
            del_cl = st.multiselect("삭제할 거래처", clients_df["병원명"].tolist(), key="med_client_del")
            if del_cl and db.button(f"🗑 {len(del_cl)}개 삭제", type="primary", use_container_width=True, key="med_client_del_btn"):
                conn = get_conn()
                for name in del_cl:
                    conn.execute("DELETE FROM med_clients WHERE id=?", (int(clients_df[clients_df["병원명"]==name].iloc[0]["id"]),))
                conn.commit(); conn.close(); st.success(f"{len(del_cl)}개 삭제 완료"); st.rerun()

    # ── 계약서 발급 ──
    with tab3:
        tmpl_row = run_sql("SELECT value FROM med_settings WHERE key='contract_template_id'")
        has_tmpl = not tmpl_row.empty and tmpl_row.iloc[0]["value"]

        if has_tmpl:
            st.success("✅ 계약서 양식 등록됨")
        else:
            st.warning("계약서 양식이 없습니다. 먼저 업로드해주세요.")

        with st.expander("📤 계약서 양식 업로드 (관리자)"):
            st.caption("양식 .docx 파일 안에 {{병원명}}, {{공급가격}} 등 플레이스홀더를 삽입해두세요.")
            tmpl_file = st.file_uploader("계약서 양식 (.docx)", type=["docx"], key="med_tmpl_up")
            if tmpl_file and st.button("양식 저장", type="primary", key="med_tmpl_save"):
                if has_tmpl:
                    drive_delete(tmpl_row.iloc[0]["value"])
                fid, _ = drive_upload(tmpl_file.read(), "_CONTRACT_TEMPLATE_.docx",
                                      "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                if fid:
                    execute("INSERT INTO med_settings (key,value) VALUES (?,?) ON CONFLICT(key) DO UPDATE SET value=EXCLUDED.value",
                            ("contract_template_id", fid))
                    st.success("양식 저장 완료!"); st.rerun()

        st.divider()
        st.markdown("**계약서 생성**")
        clients_c = run_sql("SELECT id, hospital_name, COALESCE(ceo_name,'') as ceo, COALESCE(address,'') as address FROM med_clients ORDER BY hospital_name")

        cm = st.radio("병원 입력", ["선택", "직접 입력"], horizontal=True, key="med_cont_mode")
        if cm == "선택" and not clients_c.empty:
            cont_hosp = st.selectbox("병원명", clients_c["hospital_name"].tolist(), key="med_cont_hosp")
            hr = clients_c[clients_c["hospital_name"] == cont_hosp].iloc[0]
            cont_ceo  = hr["ceo"]
            cont_addr = hr["address"]
        else:
            cont_hosp = st.text_input("병원명", key="med_cont_hosp_txt")
            cont_ceo  = st.text_input("대표자명", key="med_cont_ceo_txt")
            cont_addr = st.text_input("소재지", key="med_cont_addr_txt")

        c1, c2 = st.columns(2)
        cont_prod    = c1.text_input("제품명", key="med_cont_prod")
        cont_price   = c2.number_input("공급가격 (VAT포함)", min_value=0, step=1000000, key="med_cont_price")
        cont_pay     = st.radio("지불방법", ["현금", "리스", "일시불"], horizontal=True, key="med_cont_pay")
        cont_date    = st.date_input("계약일", value=date.today(), key="med_cont_date")

        if st.button("📄 계약서 생성 → 다운로드", type="primary", use_container_width=True, key="med_cont_gen"):
            if not has_tmpl:
                st.error("계약서 양식을 먼저 업로드해주세요.")
            elif not cont_hosp or not cont_price:
                st.error("병원명과 공급가격을 입력해주세요.")
            else:
                try:
                    from docx import Document
                    tmpl_bytes = drive_download_bytes(tmpl_row.iloc[0]["value"])
                    if tmpl_bytes is None:
                        st.error("템플릿 파일을 불러올 수 없습니다.")
                    else:
                        doc = Document(io.BytesIO(tmpl_bytes))
                        repls = {
                            "{{병원명}}":   cont_hosp,
                            "{{공급가격}}": f"₩{int(cont_price):,}원",
                            "{{제품명}}":   cont_prod,
                            "{{지불방법}}": cont_pay,
                            "{{계약일}}":   cont_date.strftime("%Y년 %m월 %d일"),
                            "{{대표자}}":   cont_ceo,
                            "{{소재지}}":   cont_addr,
                        }
                        def _replace_para(para):
                            for k, v in repls.items():
                                if k in para.text:
                                    for run in para.runs:
                                        if k in run.text:
                                            run.text = run.text.replace(k, v)
                        for para in doc.paragraphs:
                            _replace_para(para)
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for para in cell.paragraphs:
                                        _replace_para(para)
                        cbuf = io.BytesIO()
                        doc.save(cbuf); cbuf.seek(0)
                        st.download_button("⬇ 계약서 다운로드", data=cbuf,
                                           file_name=f"계약서_{cont_hosp}_{cont_date}.docx",
                                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                           use_container_width=True)
                except Exception as e:
                    st.error(f"계약서 생성 오류: {e}")

    # ── 문서함 ──
    with tab4:
        show_storage_warning()
        st.markdown("**파일 업로드**")
        clients_doc = run_sql("SELECT id, hospital_name FROM med_clients ORDER BY hospital_name")

        dc1, dc2 = st.columns(2)
        doc_hosp_mode = dc1.radio("병원", ["선택", "직접입력"], horizontal=True, key="doc_hosp_mode")
        if doc_hosp_mode == "선택" and not clients_doc.empty:
            doc_hosp = dc1.selectbox("병원명", clients_doc["hospital_name"].tolist(), key="doc_hosp_sel", label_visibility="collapsed")
            doc_cid  = int(clients_doc[clients_doc["hospital_name"] == doc_hosp].iloc[0]["id"])
        else:
            doc_hosp = dc1.text_input("병원명 직접입력", key="doc_hosp_txt", label_visibility="collapsed")
            doc_cid  = None

        doc_type = dc2.selectbox("문서 종류", ["납품계약서", "매매계약서", "기타"], key="doc_type")
        doc_file = st.file_uploader("파일 선택 (JPG / PNG / PDF)", type=["jpg","jpeg","png","pdf"], key="doc_file_up")

        if doc_file and st.button("📤 업로드", type="primary", key="doc_up_btn"):
            if not doc_hosp:
                st.error("병원명을 입력해주세요.")
            else:
                if doc_cid is None:
                    exc = run_sql("SELECT id FROM med_clients WHERE hospital_name=?", (doc_hosp,))
                    if exc.empty:
                        execute("INSERT INTO med_clients (hospital_name) VALUES (?)", (doc_hosp,))
                    doc_cid = int(run_sql("SELECT id FROM med_clients WHERE hospital_name=?", (doc_hosp,)).iloc[0]["id"])
                fname = f"{doc_hosp}_{doc_type}_{date.today()}{Path(doc_file.name).suffix}"
                fid, furl = drive_upload(doc_file.read(), fname, doc_file.type)
                if fid:
                    execute("""INSERT INTO med_contracts (client_id,file_name,file_type,drive_file_id,drive_file_url,uploaded_at)
                               VALUES (?,?,?,?,?,?)""",
                            (doc_cid, fname, doc_type, fid, furl, str(date.today())))
                    st.success(f"업로드 완료 — {fname}"); st.rerun()

        st.divider()
        st.markdown("**문서 목록**")
        flt = st.selectbox("병원 필터", ["전체"] + (clients_doc["hospital_name"].tolist() if not clients_doc.empty else []), key="doc_filter")

        docs_q = """SELECT d.id, c.hospital_name as 병원명, d.file_name as 파일명,
                           d.file_type as 종류, d.uploaded_at as 업로드일,
                           d.drive_file_url as url, d.drive_file_id as fid
                    FROM med_contracts d LEFT JOIN med_clients c ON d.client_id=c.id"""
        if flt == "전체":
            docs = run_sql(docs_q + " ORDER BY d.uploaded_at DESC")
        else:
            cid_f = int(clients_doc[clients_doc["hospital_name"] == flt].iloc[0]["id"])
            docs  = run_sql(docs_q + " WHERE d.client_id=? ORDER BY d.uploaded_at DESC", (cid_f,))

        if docs.empty:
            st.info("업로드된 문서가 없습니다.")
        else:
            for _, doc in docs.iterrows():
                dc1, dc2, dc3, dc4 = st.columns([4, 1, 1, 1])
                dc1.markdown(f"**{doc['파일명']}**  \n{doc['병원명']} | {doc['종류']} | {doc['업로드일']}")
                if dc2.button("🔗 보기", key=f"dv_{doc['id']}"):
                    st.markdown(f"[📄 Google Drive에서 열기]({doc['url']})")
                if dc3.button("⬇ 저장", key=f"dd_{doc['id']}"):
                    fb = drive_download_bytes(doc["fid"])
                    if fb:
                        ext = Path(doc["파일명"]).suffix.lower()
                        mime = "application/pdf" if ext == ".pdf" else "image/jpeg"
                        st.download_button("⬇", data=fb, file_name=doc["파일명"], mime=mime,
                                           key=f"ddb_{doc['id']}", use_container_width=True)
                if dc4.button("🗑", key=f"ddel_{doc['id']}"):
                    drive_delete(doc["fid"])
                    execute("DELETE FROM med_contracts WHERE id=?", (int(doc["id"]),))
                    st.rerun()


# ══════════════════════════════════════════════════════════════════════════
# 메인
# ══════════════════════════════════════════════════════════════════════════
init_db()

DEFAULT_MENU   = ["📊 리포트", "📦 영업 관리", "💰 수금 현황", "⚙️ 기초 관리"]
DEFAULT_LABELS = {k: k for k in DEFAULT_MENU}

if "menu_order" not in st.session_state or "menu_labels" not in st.session_state:
    _order, _labels, _site, _caption, _icon, _theme = load_menu_settings(DEFAULT_MENU, DEFAULT_LABELS)
    st.session_state.menu_order    = _order
    st.session_state.menu_labels   = _labels
    st.session_state.site_name     = _site
    st.session_state.site_caption  = _caption
    st.session_state.site_icon     = _icon
    st.session_state.theme         = _theme
if "theme" not in st.session_state:
    st.session_state.theme = "dark"
if "editing_menu" not in st.session_state:
    st.session_state.editing_menu = None

_inject_css(st.session_state.theme == "dark")

order  = st.session_state.menu_order
labels = st.session_state.menu_labels

# 헤더 — 브라우저 탭 제목 동적 업데이트 (파비콘은 set_page_config에서 처리)
st.markdown(f"""
<script>
(function() {{
    window.parent.document.title = {json.dumps(st.session_state.site_name)};
}})();
</script>
""", unsafe_allow_html=True)

hc1, hc2 = st.columns([8, 1])
with hc1:
    st.markdown(f"## {st.session_state.site_name}")
    st.caption(st.session_state.site_caption)
with hc2:
    with st.popover("⚙️ 설정", use_container_width=True):
        if "is_admin" not in st.session_state:
            st.session_state.is_admin = False

        # 테마 토글 (로그인 없이 누구나 사용)
        st.caption("테마")
        _cur_theme = st.session_state.get("theme", "dark")
        _theme_choice = st.radio(
            "테마", ["🌙 다크", "☀️ 라이트"],
            index=0 if _cur_theme == "dark" else 1,
            key="theme_radio", horizontal=True, label_visibility="collapsed"
        )
        _new_theme = "dark" if _theme_choice == "🌙 다크" else "light"
        if _new_theme != _cur_theme:
            st.session_state.theme = _new_theme
            save_menu_settings(
                st.session_state.menu_order, st.session_state.menu_labels,
                theme=_new_theme
            )
            st.rerun()
        st.divider()

        if not st.session_state.is_admin:
            st.caption("관리자 로그인")
            pw_input = st.text_input("비밀번호", type="password", key="admin_pw_input",
                                     label_visibility="collapsed", placeholder="비밀번호 입력")
            if st.button("확인", key="admin_login_btn", use_container_width=True):
                if hashlib.sha256(pw_input.encode()).hexdigest() == ADMIN_PW_HASH:
                    st.session_state.is_admin = True
                    st.rerun()
                else:
                    st.error("비밀번호가 틀렸습니다.")
        else:
            ac1, ac2 = st.columns([3, 1])
            ac1.caption("관리자 모드")
            if ac2.button("로그아웃", key="admin_logout", use_container_width=True):
                st.session_state.is_admin = False
                st.rerun()
            st.divider()
            st.caption("사이트명 변경")
            new_site_name    = st.text_input("사이트명", value=st.session_state.site_name,    key="input_site_name")
            new_site_caption = st.text_input("부제목",   value=st.session_state.site_caption, key="input_site_caption")
            if st.button("사이트명 저장", use_container_width=True, key="save_site"):
                st.session_state.site_name    = new_site_name.strip() or st.session_state.site_name
                st.session_state.site_caption = new_site_caption.strip() or st.session_state.site_caption
                save_menu_settings(
                    st.session_state.menu_order, st.session_state.menu_labels,
                    st.session_state.site_name, st.session_state.site_caption,
                    st.session_state.get("site_icon", DEFAULT_SITE_ICON)
                )
                st.rerun()
            st.divider()
            st.caption("파비콘 변경")
            _fav_path = Path(__file__).parent / "data" / "favicon.png"
            if _fav_path.exists():
                st.image(str(_fav_path), width=32, caption="현재 파비콘")
            else:
                st.caption(f"현재: {st.session_state.get('site_icon', DEFAULT_SITE_ICON)} (이모지)")
            uploaded_favicon = st.file_uploader(
                "파비콘 이미지 업로드", type=["png", "ico", "jpg", "jpeg"],
                key="favicon_uploader", label_visibility="collapsed"
            )
            fc1, fc2 = st.columns(2)
            if fc1.button("이미지 적용", key="apply_favicon", use_container_width=True, disabled=uploaded_favicon is None):
                try:
                    _fav_path.parent.mkdir(exist_ok=True)
                    _fav_path.write_bytes(uploaded_favicon.read())
                    st.success("파비콘 변경 완료!")
                    st.rerun()
                except Exception:
                    st.warning("클라우드 환경에서는 파비콘 변경이 저장되지 않습니다.")
            if fc2.button("기본 이모지로", key="reset_favicon", use_container_width=True, disabled=not _fav_path.exists()):
                _fav_path.unlink(missing_ok=True)
                st.rerun()
            st.divider()
            st.caption("탭 이름 · 순서 변경")
            for i, key in enumerate(order):
                if st.session_state.editing_menu == i:
                    new_val = st.text_input(
                        "새 이름", value=labels[key],
                        key=f"rename_input_{i}", label_visibility="collapsed"
                    )
                    sc1, sc2 = st.columns(2)
                    if sc1.button("저장", key=f"save_rename_{i}", use_container_width=True):
                        st.session_state.menu_labels[key] = new_val.strip() or labels[key]
                        st.session_state.editing_menu = None
                        save_menu_settings(st.session_state.menu_order, st.session_state.menu_labels)
                        st.rerun()
                    if sc2.button("취소", key=f"cancel_rename_{i}", use_container_width=True):
                        st.session_state.editing_menu = None
                        st.rerun()
                else:
                    c1, c2, c3, c4 = st.columns([3, 1, 1, 1])
                    c1.caption(labels[key])
                    if c2.button("✏️", key=f"me_{i}", use_container_width=True):
                        st.session_state.editing_menu = i
                        st.rerun()
                    if i > 0:
                        if c3.button("↑", key=f"mu_{i}", use_container_width=True):
                            order[i], order[i-1] = order[i-1], order[i]
                            save_menu_settings(order, st.session_state.menu_labels)
                            st.rerun()
                    if i < len(order) - 1:
                        if c4.button("↓", key=f"md_{i}", use_container_width=True):
                            order[i], order[i+1] = order[i+1], order[i]
                            save_menu_settings(order, st.session_state.menu_labels)
                            st.rerun()
            st.divider()
            if st.button("기본값으로 초기화", use_container_width=True):
                st.session_state.menu_order   = DEFAULT_MENU.copy()
                st.session_state.menu_labels  = DEFAULT_LABELS.copy()
                st.session_state.site_name    = DEFAULT_SITE_NAME
                st.session_state.site_caption = DEFAULT_SITE_CAPTION
                st.session_state.site_icon    = DEFAULT_SITE_ICON
                st.session_state.editing_menu = None
                save_menu_settings(DEFAULT_MENU, DEFAULT_LABELS, DEFAULT_SITE_NAME, DEFAULT_SITE_CAPTION, DEFAULT_SITE_ICON)
                st.rerun()

st.markdown("---")

# 대분류 선택
sector = st.radio(
    "", ["💄 화장품", "🏥 의료장비"],
    horizontal=True, key="sector_radio", label_visibility="collapsed"
)

st.markdown("---")

if sector == "💄 화장품":
    PAGE_MAP = {
        "📊 리포트":   page_dashboard,
        "📦 영업 관리": page_stock_entry,
        "💰 수금 현황": page_receivables,
        "⚙️ 기초 관리": page_master,
    }
    order = [k for k in order if k in PAGE_MAP]
    nav_tabs = st.tabs([labels[k] for k in order])
    for nav_tab, key in zip(nav_tabs, order):
        with nav_tab:
            st.title(labels[key])
            PAGE_MAP[key]()
else:
    MED_TABS = ["📊 리포트", "📦 영업 관리", "💰 수금 현황", "⚙️ 기초 관리"]
    MED_PAGE_MAP = {
        "📊 리포트":   page_med_report,
        "📦 영업 관리": page_med_sales,
        "💰 수금 현황": page_med_receivables,
        "⚙️ 기초 관리": page_med_master,
    }
    med_tabs = st.tabs(MED_TABS)
    for med_tab, key in zip(med_tabs, MED_TABS):
        with med_tab:
            st.title(key)
            MED_PAGE_MAP[key]()
